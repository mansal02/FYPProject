from __future__ import annotations

import argparse, contextlib, csv, difflib, fnmatch, importlib, io, json, mimetypes, os, platform, random, re, shutil, smtplib, subprocess, sys, threading, time, webbrowser
import pandas as pd
from email.message import EmailMessage
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import quote_plus
import win32com.client as win32
from datetime import datetime, timedelta

from openpyxl import Workbook, load_workbook
from openpyxl.chart import LineChart, Reference
from openpyxl.utils import get_column_letter

from aiassistant.infra.config.app_config import CONFIG
from aiassistant.infra.db.database_manager import DatabaseManager
# =============================================================================
# 1. OPTIONAL DEPENDENCIES
# =============================================================================

try:
    import pyautogui
    pyautogui.FAILSAFE = True
except ImportError:
    pyautogui = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
    Document = docx.Document
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    Document = None
    DOCX_AVAILABLE = False

try:
    import numpy as np
except ImportError:
    np = None

try:
    from duckduckgo_search import DDGS
except ImportError:
    DDGS = None

try:
    import win32com.client as win32
except ImportError:
    win32 = None

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    requests = None
    REQUESTS_AVAILABLE = False

try:
    import pywhatkit
except ImportError:
    pywhatkit = None

try:
    from send2trash import send2trash
except ImportError:
    send2trash = None

try:
    from AppOpener import close as appopener_close, give_appnames as appopener_list, open as appopener_open
    open_app, close_app, give_appnames = appopener_open, appopener_close, appopener_list
    APPOPENER_AVAILABLE = True
except ImportError:
    appopener_open, appopener_close, appopener_list = None, None, None
    open_app, close_app, give_appnames = None, None, None
    APPOPENER_AVAILABLE = False

try:
    BeautifulSoup = importlib.import_module("bs4").BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BeautifulSoup = None
    BS4_AVAILABLE = False

try:
    pyperclip = importlib.import_module("pyperclip")
    CLIPBOARD_AVAILABLE = True
except ImportError:
    pyperclip = None
    CLIPBOARD_AVAILABLE = False

try:
    Presentation = importlib.import_module("pptx").Presentation
    POWERPOINT_AVAILABLE = True
except ImportError:
    Presentation = None
    POWERPOINT_AVAILABLE = False


# =============================================================================
# 2. CONSTANTS & CONFIGURATIONS
# =============================================================================

_SENTENCE_TRANSFORMER_CLASS = None
_SENTENCE_TRANSFORMER_UNAVAILABLE = False

_SAFE_MODE_ENABLED = bool(CONFIG.get("actions", {}).get("safe_mode", True))
_PROTECTED_PATH_PREFIXES = []

SYSTEM_NAME = platform.system().lower()
IS_WINDOWS = SYSTEM_NAME.startswith("win")
TOOL_BRIDGE_VERBOSE = str(os.environ.get("MARIE_TOOL_BRIDGE_VERBOSE", "1")).strip().lower() in {"1", "true", "yes", "on"}

TEXT_SEARCH_EXTENSIONS = {
    ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".ini", ".cfg", ".csv", ".log",
    ".xml", ".html", ".htm", ".js", ".jsx", ".ts", ".tsx", ".java", ".cs", ".cpp",
    ".c", ".h", ".hpp", ".ps1", ".bat", ".cmd", ".sql", ".toml", ".rtf",
}

WINDOWS_APP_ALIASES = {
    "vscode": "code", "vs code": "code", "visual studio code": "code", "visual studio": "devenv",
    "excel": "excel", "word": "winword", "powerpoint": "powerpnt", "outlook": "outlook",
    "notepad": "notepad", "calculator": "calc", "cmd": "cmd", "music": "ytmusic", "yt": "youtube",
    "steam": r"C:\Program Files (x86)\Steam\steam.exe",
    "obs": r"C:\Program Files\obs-studio\bin\64bit\obs64.exe"
}

WEB_APP_ALIASES = {
    "youtube": "https://www.youtube.com", "yt": "https://www.youtube.com",
    "youtube music": "https://music.youtube.com", "youtube studio": "https://studio.youtube.com",
    "google": "https://www.google.com", "github": "https://github.com",
    "gmail": "https://mail.google.com", "whatsapp": "https://web.whatsapp.com",
    "telegram": "https://web.telegram.org",
}

SMART_SEARCH_STOP_WORDS = {
    "a", "an", "and", "are", "can", "did", "do", "for", "from", "get", "have", "i", "in", "is",
    "it", "last", "me", "my", "of", "on", "open", "please", "recent", "search", "that", "the",
    "this", "to", "want", "with",
}

SMART_SEARCH_SKIP_DIR_NAMES = {
    "$recycle.bin", "system volume information", ".git", ".venv", "node_modules",
    "__pycache__", "models", "piper", "rvc", "checkpoints", "chroma", "cache",
    "pkgconfig", "appdata", "program files", "program files (x86)", "windows"
}

# --- Action Command Definitions ---
GENERAL_TASK_COMMANDS = [
    "open <app>", "close <app>", "play <topic>", "write <text>", "note <text>",
    "type <text>", "take a note <text>", "volume up", "volume down", "mute", "unmute",
    "search web <query>", "open website <url>", "browse <url>", "research <query>",
    "web research <query>", "research web <query>",
]

CLIPBOARD_COMMANDS = [
    "copy selected text", "copy now", "paste clipboard", "paste now",
    "save clipboard to rad as <key>",
]

SYSTEM_COMMANDS = [
    "system check", "check system", "malware scan", "scan for malware",
    "run malware scan", "security quick scan", "scan apps", "update apps",
]

FILE_SYSTEM_COMMANDS = [
    "files roots", "files list <path>", "files deep search <query>",
    "files deep search <query> [in <root>]", "search file <hint>", "find file <hint>",
    "locate file <hint>", "files analyze <path>", "files create file <path> [content <text>]",
    "files create folder <path>", "files create dir <path>", "files create directory <path>",
    "files move <source> -> <destination>", "files copy <source> -> <destination>",
    "files delete <path>", "files remove <path>", "files open <path>", "open file <hint>",
]

SOFTWARE_COMMANDS = [
    "software open <app>", "software close <app>", "software running",
    "service open <gmail|outlook|whatsapp|telegram>", "service open <app_or_service>",
]

COMMUNICATION_COMMANDS = [
    "email draft to <email> subject <subject> body <body> [attach <file_hint>] [provider gmail|outlook|custom]",
    "email send to <email> subject <subject> body <body> [attach <file_hint>] [provider gmail|outlook|custom]",
    "telegram send to <chat_id> token <bot_token> message <text>",
    "telegram file <path_or_hint> to <chat_id> token <bot_token> [caption <text>]",
    "whatsapp send to <phone_number> message <text> (WhatsApp Web mode)",
]

ASSISTANT_TOOL_ACTIONS = [
    "list_system_roots", "list_directory", "deep_search (alias: deep_search_paths)",
    "analyze_path", "create_path", "move_path", "copy_path", "delete_path",
    "search_file", "semantic_search_file", "read_file", "open_file (alias: open_path)",
    "launch_application", "close_application", "list_running_apps (alias: list_running_applications)",
    "open_service", "search_mirror", "move_mouse", "click", "type_text", "press_key",
    "hotkey", "run_command", "toggle_dark_mode", "send_email", "draft_email_attachment",
    "send_telegram", "send_whatsapp", "online_query",
]

ASSISTANT_TOOL_EXAMPLES = [
    '<tool>{"action":"list_directory","path":"D:/pylearn/FYP/AiAssistant","max_results":120}</tool>',
    '<tool>{"action":"deep_search","query":"runsys.py","roots":["D:/pylearn/FYP/AiAssistant"],"include_content":false}</tool>',
    '<tool>{"action":"send_email","to":"name@example.com","subject":"Report","body":"Please find attached","provider":"outlook","attachments":["report.pdf"],"send_now":false}</tool>',
]

OFFICE_HELP_COMMANDS = ["office help", "excel help", "word help", "powerpoint help"]

EXCEL_HELP_COMMANDS = [
    "create an excel workbook with <rows>x<cols> table with random data and then create the graph from it",
    "excel random table <rows>x<cols> with graph [in <file>]", "excel demo table graph [in <file>]",
    "excel create <file>", "excel create sheet <sheet_name> in <file>", "excel list sheets in <file>",
    "excel set <cell> to <value> in <file> [sheet <sheet_name>]", "excel get <cell> in <file> [sheet <sheet_name>]",
    "excel add row <comma-separated values> in <file> [sheet <sheet_name>]",
    "excel delete row <number> in <file> [sheet <sheet_name>]", "excel delete column <A..Z> in <file> [sheet <sheet_name>]",
    "excel sum <A1:B10> in <file> to <cell> [sheet <sheet_name>]", "excel formula <cell> = <formula> in <file> [sheet <sheet_name>]",
]

WORD_HELP_COMMANDS = [
    "word create <file>", "word add heading <text> [level 1-6] in <file>",
    "word add paragraph <text> in <file>", "word read <file>",
]

POWERPOINT_HELP_COMMANDS = [
    "powerpoint create <file>", "powerpoint add slide title <title> content <content> in <file>",
    "powerpoint launch <file>",
]

ASSISTANT_JSON_ACTIONS = [
    '{"action":"open","target":"chrome"}',
    '{"action":"close","target":"notepad"}',
    '{"action":"search_web","target":"latest ai news"}',
    '{"action":"open_website","target":"github.com"}',
    '{"action":"volume","target":"up"}',
    '{"action":"write_note","target":"meeting summary"}',
    '{"action":"play","target":"lofi coding music"}',
]


# =============================================================================
# 3. HELPER UTILITIES
# =============================================================================

def _get_sentence_transformer_class():
    global _SENTENCE_TRANSFORMER_CLASS, _SENTENCE_TRANSFORMER_UNAVAILABLE
    if _SENTENCE_TRANSFORMER_CLASS is not None:
        return _SENTENCE_TRANSFORMER_CLASS
    if _SENTENCE_TRANSFORMER_UNAVAILABLE:
        return None
    try:
        from sentence_transformers import SentenceTransformer as _SentenceTransformer
        _SENTENCE_TRANSFORMER_CLASS = _SentenceTransformer
        return _SENTENCE_TRANSFORMER_CLASS
    except ImportError:
        _SENTENCE_TRANSFORMER_UNAVAILABLE = True
        return None

def _ok(message: str, data: Any = None) -> Dict[str, Any]:
    return {"success": True, "message": message, "data": data}

def _fail(message: str, error: str) -> Dict[str, Any]:
    return {"success": False, "message": message, "error": error}

def _to_bool(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool): return value
    if value is None: return default
    if isinstance(value, (int, float)): return bool(value)
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on"}: return True
    if text in {"0", "false", "no", "n", "off", ""}: return False
    return default

def _to_str_list(value: Any) -> List[str]:
    if value is None: return []
    if isinstance(value, list): return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    if not text: return []
    if "," in text: return [part.strip() for part in text.split(",") if part.strip()]
    return [text]

def _safe_resolve_path(raw_path: str) -> Path:
    path = Path(str(raw_path or "").strip()).expanduser()
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()
    else:
        path = path.resolve()
    return path

def _is_protected_path(path_obj: Path) -> bool:
    if not _SAFE_MODE_ENABLED: return False
    try:
        resolved = str(path_obj.resolve()).replace("\\", "/").lower()
    except Exception:
        resolved = str(path_obj).replace("\\", "/").lower()
    return any(resolved.startswith(prefix) for prefix in _PROTECTED_PATH_PREFIXES)

def _list_drive_roots() -> List[Path]:
    if os.name != "nt": return [Path("/")]
    roots: List[Path] = []
    for letter in "ABCDEFGHIJKLMNOPQRSTUVWXYZ":
        drive = Path(f"{letter}:/")
        try:
            if drive.exists(): roots.append(drive)
        except Exception:
            continue
    if not roots:
        roots.append(Path("C:/"))
    return roots

def default_search_roots() -> List[Path]:
    cwd = Path.cwd()
    roots = [cwd]
    if os.name == "nt":
        for drive in ("C:/", "D:/"):
            drive_path = Path(drive)
            try:
                if drive_path.exists(): roots.append(drive_path)
            except Exception:
                continue
    try:
        if cwd.resolve() == Path.home().resolve():
            roots.extend([Path.home() / "Desktop", Path.home() / "Documents", Path.home() / "Downloads"])
    except Exception:
        pass

    unique_existing: List[Path] = []
    seen = set()
    for item in roots:
        resolved = item.resolve()
        if resolved.exists() and resolved not in seen:
            unique_existing.append(resolved)
            seen.add(resolved)
    return unique_existing

def _open_with_default_app(path):
    try:
        if IS_WINDOWS and hasattr(os, "startfile"):
            os.startfile(path)
        elif SYSTEM_NAME == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
        return True
    except Exception as e:
        print(f"[ACTION] Could not open file with default app: {e}")
        return False

def _normalize_office_path(file_name: str, default_extension: str) -> str:
    """Unified path normalizer for Excel, Word, and PowerPoint handlers."""
    clean_name = file_name.strip().strip('"').strip("'")
    if not clean_name.lower().endswith(default_extension):
        clean_name += default_extension

    is_windows_absolute = re.match(r"^[a-zA-Z]:\\", clean_name) is not None
    if not os.path.isabs(clean_name) and not is_windows_absolute:
        clean_name = os.path.abspath(clean_name)

    parent = os.path.dirname(clean_name)
    if parent:
        os.makedirs(parent, exist_ok=True)
    return clean_name


# =============================================================================
# 4. FILE & DATA PROCESSORS
# =============================================================================

def create_outlook_meeting(subject: str, start: str, duration_minutes: int, attendees: list) -> dict:
    """Instantly returns to the LLM while Outlook opens and injects Teams in the background."""
    
    def run_outlook():
        import pythoncom
        pythoncom.CoInitialize() # Required for background thread COM access
        try:
            outlook = win32.Dispatch("Outlook.Application")
            appt = outlook.CreateItem(1) # 1 = olAppointmentItem
            
            try:
                start_time = datetime.fromisoformat(start)
            except ValueError:
                start_time = datetime.now() + timedelta(hours=1)

            appt.Start = start_time.strftime("%Y-%m-%d %H:%M")
            appt.Duration = int(duration_minutes)
            appt.Subject = subject
            appt.MeetingStatus = 1 # 1 = olMeeting

            for attendee in attendees:
                appt.Recipients.Add(attendee)
            
            # 1. We MUST display the window first so the Ribbon UI loads into memory
            appt.Display() 
            
            # 2. Give Outlook a tiny delay to render the window
            time.sleep(0.5) 
            
            # 3. Force the execution of the Teams Add-in button
            try:
                appt.GetInspector.CommandBars.ExecuteMso("AddOnlineMeeting")
                print("[Action] Teams meeting link successfully injected.")
            except Exception as ribbon_err:
                print(f"[Warning] Could not auto-click Teams button. The Add-in might be slow: {ribbon_err}")
                
        except Exception as e:
            print(f"[Outlook Error] Failed to create meeting: {str(e)}")
        finally:
            pythoncom.CoUninitialize()

    # Fire and forget
    threading.Thread(target=run_outlook, daemon=True).start()
    
    return {"success": True, "message": "Triggered Outlook in the background and requested a Teams link."}

def convert_file_to_json(file_path: str, max_chars: int = 15000) -> dict:
    """Converts supported files (.xlsx, .docx, .pdf) into a JSON structure."""
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        return {"error": "File not found"}

    ext = path.suffix.lower()
    file_data = {"filename": path.name, "type": ext, "content": None}

    try:
        # Excel/CSV
        if ext in {".xlsx", ".xls", ".csv"}:
            if ext == ".csv":
                df = pd.read_csv(path)
            else:
                df = pd.read_excel(path, sheet_name=None)
                
            if isinstance(df, dict): # Multiple sheets
                sheet_data = {}
                for sheet_name, sheet_df in df.items():
                    sheet_data[sheet_name] = json.loads(sheet_df.to_json(orient="records"))
                file_data["content"] = sheet_data
            else:
                file_data["content"] = json.loads(df.to_json(orient="records"))

        # PDF
        elif ext == ".pdf":
            if PyPDF2 is None:
                return {"success": False, "error": "PyPDF2 is missing."}
            with path.open("rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text: pages.append({"page": i+1, "text": text.strip()})
                file_data["content"] = pages

        # Word
        elif ext in {".doc", ".docx"}:
            if not DOCX_AVAILABLE:
                return {"success": False, "error": "python-docx is missing."}
            doc = Document(str(path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            file_data["content"] = {"paragraphs": paragraphs}

        else:
            # Fallback to plain text
            text = path.read_text(encoding="utf-8", errors="ignore")
            file_data["content"] = text

        # Truncate to prevent context window overflow
        json_str = json.dumps(file_data)
        if len(json_str) > max_chars:
            file_data["content"] = "Content truncated due to length. " + json_str[:max_chars]
            
        return {"success": True, "data": file_data}

    except Exception as e:
        return {"success": False, "error": str(e)}

def extract_document_text(file_path: str, instruction: str) -> dict:
    """Extracts text quickly and feeds it back to the main LLM to avoid double-loading models."""
    from aiassistant.tools.tools_os import convert_file_to_json
    
    extraction = convert_file_to_json(file_path)
    if not extraction.get("success"):
        return extraction
        
    # Return raw text directly to the agent's context window
    # The agent will read this and formulate the final answer itself.
    return {"success": True, "data": extraction["data"]["content"][:6000]} # Cap length to fit context

def execute_office_tool(tool_name: str, payload: dict) -> dict:
    """
    Dynamically routes Office Automation tasks, generates Word, PDF, or Excel files,
    and opens them on the user's screen.
    """
    try:
        save_dir = Path.home() / "Desktop" / "MARIE_Reports"
        save_dir.mkdir(parents=True, exist_ok=True)
        
        # --- WORD EXPORT ---
        if tool_name == "create_word_doc":
            if not DOCX_AVAILABLE:
                return {"success": False, "error": "The 'python-docx' library is missing."}
            
            filename = payload.get("filename", "MARIE_Report.docx")
            if not filename.endswith(".docx"): filename += ".docx"
            filepath = save_dir / filename

            doc = Document()
            doc.add_heading(payload.get("title", "Data Analysis Report"), 0)
            
            for paragraph_text in payload.get("content", "").split('\n'):
                if paragraph_text.strip(): doc.add_paragraph(paragraph_text.strip())
            
            doc.save(str(filepath))
            if hasattr(os, 'startfile'): os.startfile(str(filepath))
            return {"success": True, "message": f"Word report generated at {filepath}"}
            
        # --- PDF EXPORT ---
        elif tool_name == "create_pdf_report":
            try:
                from fpdf import FPDF
            except ImportError:
                return {"success": False, "error": "The 'fpdf' library is missing. Run: pip install fpdf"}

            filename = payload.get("filename", "MARIE_Report.pdf")
            if not filename.endswith(".pdf"): filename += ".pdf"
            filepath = save_dir / filename

            pdf = FPDF()
            pdf.add_page()
            
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, txt=payload.get("title", "Data Analysis Report"), ln=True, align='C')
            
            pdf.set_font("Arial", size=12)
            pdf.ln(10)
            
            content = payload.get("content", "No analysis content provided.")
            clean_content = content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, txt=clean_content)
            
            pdf.output(str(filepath))
            if hasattr(os, 'startfile'): os.startfile(str(filepath))
            return {"success": True, "message": f"PDF report generated at {filepath}"}

        # --- EXCEL EXPORT ---
        elif tool_name == "create_excel_report":
            if 'openpyxl' not in globals() and 'openpyxl' not in sys.modules:
                return {"success": False, "error": "The 'openpyxl' library is missing."}
                
            filename = payload.get("filename", "MARIE_Data.xlsx")
            if not filename.endswith(".xlsx"): filename += ".xlsx"
            filepath = save_dir / filename

            wb = Workbook()
            ws = wb.active
            ws.title = "Analysis Report"

            table_data = payload.get("table_data", [])
            if not table_data:
                ws.append(["Notice", "No structured data provided for Excel export."])
            else:
                for row in table_data:
                    if isinstance(row, dict):
                        if ws.max_row == 1 and not ws.cell(1,1).value:
                            ws.append(list(row.keys()))
                        ws.append(list(row.values()))
                    elif isinstance(row, list):
                        ws.append(row)
                    else:
                        ws.append([str(row)])

            wb.save(str(filepath))
            if hasattr(os, 'startfile'): os.startfile(str(filepath))
            return {"success": True, "message": f"Excel report generated at {filepath}"}
            
        else:
            return {"success": False, "error": f"Unknown tool requested: {tool_name}"}
            
    except PermissionError:
        return {"success": False, "error": "Permission denied. Please close the file if it is currently open."}
    except Exception as e:
        return {"success": False, "error": f"Tool execution crashed: {str(e)}"}


# =============================================================================
# 5. CORE CLASSES & UNIFIED TOOL BRIDGE
# =============================================================================

class ConnectivityModule:
    """Online query module backed by duckduckgo-search."""
    def online_query(self, query: str, max_results: int = 5) -> Dict[str, Any]:
        clean = (query or "").strip()
        if not clean:
            return _fail("Online query text was empty.", "empty_online_query")

        if DDGS is None:
            return _fail("duckduckgo-search is unavailable.", "missing_dependency: duckduckgo-search")

        try:
            with DDGS() as ddgs:
                results = list(ddgs.text(clean, max_results=max(1, int(max_results))))
        except Exception as exc:
            return _fail("Online query failed.", str(exc))

        compact = []
        for item in results:
            if not isinstance(item, dict): continue
            compact.append({
                "title": str(item.get("title", "")).strip(),
                "url": str(item.get("href", "")).strip(),
                "snippet": str(item.get("body", "")).strip(),
            })

        return _ok(f"Found {len(compact)} online result(s).", data=compact)


class UnifiedTaskBridge:
    """Unified bridge for UI automation, semantic file search, and system commands."""

    def __init__(self) -> None:
        self.connectivity = ConnectivityModule()
        self._semantic_model = None

    @staticmethod
    def _is_vague_hint(hint: str) -> bool:
        clean = (hint or "").strip().lower()
        if not clean: return False
        if len(clean.split()) >= 3: return True
        vague_tokens = {"something", "document", "notes", "file", "thing", "report", "project", "old", "latest"}
        return any(token in clean for token in vague_tokens)

    def _load_semantic_model(self):
        if np is None: return None
        sentence_transformer_cls = _get_sentence_transformer_class()
        if sentence_transformer_cls is None: return None
        if self._semantic_model is None:
            try:
                self._semantic_model = sentence_transformer_cls("all-MiniLM-L6-v2", device="cpu")
            except Exception:
                return None
        return self._semantic_model

    @staticmethod
    def _iter_candidate_files(search_roots: List[Path], max_files: int = 1400) -> List[Path]:
        files: List[Path] = []
        priority_dirs = {
            "aiassistant": 0, "knowledge": 1, "script": 2, "launchers": 3,
            "frontend": 4, "backend": 5, "core": 6, "tools": 7, "infra": 8,
        }

        for root in search_roots:
            try:
                if not root.exists(): continue
                for current_root, dirs, filenames in os.walk(root, topdown=True, onerror=lambda _e: None):
                    dirs[:] = [d for d in dirs if d.lower() not in SMART_SEARCH_SKIP_DIR_NAMES]
                    dirs.sort(key=lambda d: (priority_dirs.get(d.lower(), 99), d.lower()))

                    current_path = Path(current_root)
                    for filename in filenames:
                        path = current_path / filename
                        if len(files) >= max_files: return files
                        if not path.is_file(): continue
                        try:
                            if path.stat().st_size > 20 * 1024 * 1024: continue
                        except Exception:
                            continue
                        files.append(path)
            except Exception:
                continue
        return files

    @staticmethod
    def _path_semantic_text(path: Path) -> str:
        stem_tokens = re.sub(r"[_\-\.]+", " ", path.stem)
        parent_tokens = re.sub(r"[_\-\.]+", " ", str(path.parent).replace("\\", " "))
        return f"{path.name} {stem_tokens} {parent_tokens}"

    @staticmethod
    def _normalize_search_text(text: str) -> str:
        return re.sub(r"[^a-z0-9]+", " ", str(text or "").lower()).strip()

    @classmethod
    def _tokenize_search_hint(cls, hint: str) -> List[str]:
        tokens: List[str] = []
        seen = set()
        for token in re.findall(r"[a-z0-9]{2,}", cls._normalize_search_text(hint)):
            if token in SMART_SEARCH_STOP_WORDS: continue
            if token not in seen:
                tokens.append(token)
                seen.add(token)
        return tokens

    @classmethod
    def _score_path_hint_match(cls, hint: str, hint_tokens: List[str], path: Path) -> float:
        hint_norm = cls._normalize_search_text(hint)
        name_norm = cls._normalize_search_text(path.name)
        stem_norm = cls._normalize_search_text(path.stem)
        full_norm = cls._normalize_search_text(str(path))

        if not hint_norm: return 0.0

        ratio_name = difflib.SequenceMatcher(None, hint_norm, name_norm).ratio()
        ratio_stem = difflib.SequenceMatcher(None, hint_norm, stem_norm).ratio()
        ratio_full = difflib.SequenceMatcher(None, hint_norm, full_norm).ratio()

        token_hits = sum(1 for token in hint_tokens if token in full_norm)
        token_score = (token_hits / len(hint_tokens)) if hint_tokens else 0.0

        exact_bonus = 0.0
        if hint_norm in name_norm: exact_bonus = 0.45
        elif hint_norm in full_norm: exact_bonus = 0.30

        weighted = max((ratio_name * 0.55) + (ratio_stem * 0.25) + (ratio_full * 0.20), token_score)
        return min(1.0, weighted + exact_bonus)

    @classmethod
    def _content_grep_score(cls, path: Path, hint_tokens: List[str], hint: str, max_file_size_bytes: int = 3 * 1024 * 1024) -> tuple[float, str]:
        if not hint_tokens: return 0.0, ""
        if path.suffix.lower() not in TEXT_SEARCH_EXTENSIONS: return 0.0, ""
        try:
            if path.stat().st_size > max_file_size_bytes: return 0.0, ""
        except Exception:
            return 0.0, ""

        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return 0.0, ""

        if not text: return 0.0, ""

        hint_norm = cls._normalize_search_text(hint)
        best_score = 0.0
        best_line = ""

        for raw_line in text.splitlines()[:1800]:
            line = raw_line.strip()
            if not line: continue

            line_norm = cls._normalize_search_text(line)
            if not line_norm: continue

            token_hits = sum(1 for token in hint_tokens if token in line_norm)
            if token_hits <= 0: continue

            coverage = token_hits / len(hint_tokens)
            fuzzy = difflib.SequenceMatcher(None, hint_norm, line_norm).ratio() if hint_norm else 0.0
            score = max(coverage, fuzzy * 0.85)

            if score > best_score:
                best_score = score
                best_line = line[:220]
                if score >= 0.95: break

        return best_score, best_line

    def smart_heuristic_search_files(self, hint: str, roots: Optional[List[str]] = None, max_results: int = 20, include_content: bool = True, max_scan_files: int = 4200) -> Dict[str, Any]:
        clean_hint = str(hint or "").strip()
        if not clean_hint: return _fail("Search hint was empty.", "empty_search_hint")
        
        cap_results = max(1, int(max_results))
        
        try:
            db = DatabaseManager()
            if hasattr(db, "search_searchable_mirror"):
                mirror_results = db.search_searchable_mirror(clean_hint, limit=cap_results)
                if mirror_results:
                    valid_paths = []
                    for entry in mirror_results:
                        file_path_str = entry.get("file_path")
                        if file_path_str and os.path.exists(file_path_str):
                            valid_paths.append(str(Path(file_path_str).resolve()))
                    
                    if valid_paths:
                        return _ok(f"Fast index mirror search located {len(valid_paths)} match(es).", data=valid_paths[:cap_results])
        except Exception as db_err:
            print(f"[Warning] Fast searchable mirror index lookup bypassed: {db_err}")

        search_roots = self._resolve_search_roots(roots=roots, include_all_drives=False)
        if not search_roots: return _fail("No accessible roots were found for search.", "missing_roots")

        files = self._iter_candidate_files(search_roots, max_files=max(300, int(max_scan_files)))
        if not files: return _ok("No files available for smart heuristic search.", data=[])

        hint_tokens = self._tokenize_search_hint(clean_hint)
        path_ranked: List[Dict[str, Any]] = []
        
        for file_path in files:
            try:
                if any(part.lower() in SMART_SEARCH_SKIP_DIR_NAMES for part in file_path.parts): continue
            except Exception:
                pass

            path_score = self._score_path_hint_match(clean_hint, hint_tokens, file_path)
            if path_score < 0.18 and hint_tokens: continue

            path_ranked.append({"path": file_path, "path_score": path_score, "content_score": 0.0, "content_line": ""})

        if not path_ranked:
            return self.lexical_search_files(clean_hint, roots=roots, max_results=cap_results)

        path_ranked.sort(key=lambda item: float(item["path_score"]), reverse=True)
        shortlist_cap = min(max(len(path_ranked), cap_results * 8), 320)
        shortlist = path_ranked[:shortlist_cap]

        if include_content and hint_tokens:
            for item in shortlist:
                try:
                    score, line = self._content_grep_score(path=item["path"], hint_tokens=hint_tokens, hint=clean_hint)
                    item["content_score"] = score
                    item["content_line"] = line
                except Exception:
                    continue

        scored: List[Dict[str, Any]] = []
        for item in shortlist:
            combined = (float(item["path_score"]) * 0.72) + (float(item["content_score"]) * 0.28)
            if combined < 0.20 and float(item["path_score"]) < 0.28: continue
            try:
                scored.append({"path": str(item["path"].resolve()), "score": round(combined, 4), "snippet": str(item["content_line"])})
            except Exception:
                continue

        if not scored: return self.lexical_search_files(clean_hint, roots=roots, max_results=cap_results)

        scored.sort(key=lambda item: float(item["score"]), reverse=True)
        result_paths = [entry["path"] for entry in scored[:cap_results]]
        return _ok(f"Smart heuristic search found {len(result_paths)} match(es).", data=result_paths)

    def lexical_search_files(self, hint: str, roots: Optional[List[str]] = None, max_results: int = 20) -> Dict[str, Any]:
        clean = (hint or "").strip()
        if not clean: return _fail("Search hint was empty.", "empty_search_hint")

        hint_lower = clean.lower()
        search_roots = [Path(r) for r in roots] if roots else default_search_roots()

        matches: List[str] = []
        for root in search_roots:
            try:
                if not root.exists(): continue
                for path in root.rglob("*"):
                    if len(matches) >= max_results: break
                    if not path.is_file(): continue
                    if hint_lower in path.name.lower(): matches.append(str(path.resolve()))
                if len(matches) >= max_results: break
            except Exception:
                continue

        return _ok(f"Found {len(matches)} lexical file match(es).", data=matches)

    def semantic_search_files(self, hint: str, roots: Optional[List[str]] = None, max_results: int = 20) -> Dict[str, Any]:
        clean = (hint or "").strip()
        if not clean: return _fail("Search hint was empty.", "empty_search_hint")

        model = self._load_semantic_model()
        if model is None:
            return self.lexical_search_files(clean, roots=roots, max_results=max_results)

        search_roots = [Path(r) for r in roots] if roots else default_search_roots()
        files = self._iter_candidate_files(search_roots)
        if not files: return _ok("No files available for semantic search.", data=[])

        corpus = [self._path_semantic_text(path) for path in files]

        try:
            query_vec = model.encode([clean], normalize_embeddings=True)
            corpus_vec = model.encode(corpus, normalize_embeddings=True)
            scores = np.dot(corpus_vec, query_vec[0])
            ranked_idx = np.argsort(scores)[::-1][: max(1, int(max_results))]
        except Exception as exc:
            return _fail("Semantic file search failed.", str(exc))

        results: List[str] = [str(files[int(idx)].resolve()) for idx in ranked_idx]
        return _ok(f"Found {len(results)} semantic file match(es).", data=results)

    def search_files_by_hint(self, hint: str, roots: Optional[List[str]] = None, max_results: int = 20) -> Dict[str, Any]:
        heuristic = self.smart_heuristic_search_files(hint=hint, roots=roots, max_results=max_results, include_content=True)
        if isinstance(heuristic, dict) and heuristic.get("success") and heuristic.get("data"):
            return heuristic

        if self._is_vague_hint(hint):
            semantic = self.semantic_search_files(hint, roots=roots, max_results=max_results)
            if isinstance(semantic, dict) and semantic.get("success") and semantic.get("data"):
                return semantic

        return self.lexical_search_files(hint, roots=roots, max_results=max_results)

    def _resolve_search_roots(self, roots: Optional[List[str]] = None, include_all_drives: bool = False) -> List[Path]:
        candidates: List[Path] = []
        if roots:
            for root in roots:
                clean = str(root or "").strip()
                if not clean: continue
                try: candidates.append(_safe_resolve_path(clean))
                except Exception: continue
        else:
            candidates.extend(default_search_roots())

        if include_all_drives:
            candidates.extend(_list_drive_roots())

        unique_existing: List[Path] = []
        seen = set()
        for item in candidates:
            try: resolved = item.resolve()
            except Exception: continue
            if resolved.exists() and resolved not in seen:
                unique_existing.append(resolved)
                seen.add(resolved)
        return unique_existing

    def list_system_roots(self) -> Dict[str, Any]:
        roots = self._resolve_search_roots(include_all_drives=True)
        if not roots: return _fail("No accessible roots were found.", "missing_roots")
        return _ok("Listed accessible storage roots.", data=[str(path) for path in roots])

    @staticmethod
    def _mirror_requested(action: Dict[str, Any]) -> bool:
        return _to_bool(action.get("use_mirror", False), default=False) or bool(str(action.get("mirror_query") or "").strip())

    @staticmethod
    def _mirror_query_from_action(action: Dict[str, Any], fallback: str) -> str:
        raw = action.get("mirror_query") or action.get("mirror_hint") or action.get("mirror_text")
        if raw is None: raw = fallback
        return str(raw or "").strip()

    def _get_mirror_context(self, query: str, limit: int = 6) -> Dict[str, Any] | None:
        clean = str(query or "").strip()
        if not clean: return None
        try:
            db = DatabaseManager()
            results = db.search_searchable_mirror(clean, limit=max(1, int(limit)))
        except Exception:
            return None
        if not results: return None
        return {"query": clean, "matches": results}

    @staticmethod
    def _attach_mirror_context(result: Dict[str, Any], mirror_context: Dict[str, Any] | None) -> Dict[str, Any]:
        if not mirror_context or not isinstance(result, dict): return result
        result["mirror_context"] = mirror_context
        return result

    def search_mirror(self, query: str, limit: int = 8) -> Dict[str, Any]:
        clean = str(query or "").strip()
        if not clean: return _fail("Mirror search query was empty.", "empty_query")
        try:
            db = DatabaseManager()
            results = db.search_searchable_mirror(clean, limit=int(limit) or 8)
            return _ok(f"Mirror search returned {len(results)} match(es).", data=results)
        except Exception as exc:
            return _fail("Mirror search failed.", str(exc))

    def list_directory(self, path: str, recursive: bool = False, max_depth: int = 2, max_results: int = 300, pattern: str = "", include_files: bool = True, include_dirs: bool = True) -> Dict[str, Any]:
        clean_path = str(path or "").strip()
        if not clean_path: return _fail("Directory path was empty.", "empty_path")

        try:
            target = _safe_resolve_path(clean_path)
            if not target.exists() or not target.is_dir():
                return _fail("Directory was not found.", f"missing_directory: {target}")

            max_items = max(1, int(max_results))
            depth_limit = max(0, int(max_depth))
            pattern_text = str(pattern or "").strip().lower()

            def _name_match(name: str) -> bool:
                if not pattern_text: return True
                lowered = name.lower()
                return pattern_text in lowered or fnmatch.fnmatch(lowered, pattern_text)

            items: List[Dict[str, Any]] = []
            truncated = False

            if not recursive:
                for child in sorted(target.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
                    if child.is_dir() and not include_dirs: continue
                    if child.is_file() and not include_files: continue
                    if not _name_match(child.name): continue

                    entry: Dict[str, Any] = {"path": str(child.resolve()), "type": "directory" if child.is_dir() else "file"}
                    if child.is_file():
                        try: entry["size_bytes"] = child.stat().st_size
                        except Exception: pass
                    items.append(entry)
                    if len(items) >= max_items:
                        truncated = True
                        break
            else:
                base_depth = len(target.parts)
                for current_root, dirs, files in os.walk(target, topdown=True, onerror=lambda _e: None):
                    current = Path(current_root)
                    current_depth = len(current.parts) - base_depth
                    if current_depth >= depth_limit: dirs[:] = []

                    if include_dirs:
                        for dirname in sorted(dirs):
                            if not _name_match(dirname): continue
                            items.append({"path": str((current / dirname).resolve()), "type": "directory"})
                            if len(items) >= max_items:
                                truncated = True
                                break
                        if truncated: break

                    if include_files:
                        for filename in sorted(files):
                            if not _name_match(filename): continue
                            file_path = current / filename
                            entry = {"path": str(file_path.resolve()), "type": "file"}
                            try: entry["size_bytes"] = file_path.stat().st_size
                            except Exception: pass
                            items.append(entry)
                            if len(items) >= max_items:
                                truncated = True
                                break
                        if truncated: break

            return _ok(f"Listed {len(items)} item(s) from {target}.", data={"path": str(target), "items": items, "truncated": truncated, "recursive": bool(recursive)})
        except Exception as exc:
            return _fail("Directory listing failed.", str(exc))

    def deep_search_paths(self, query: str, roots: Optional[List[str]] = None, max_results: int = 40, include_all_drives: bool = False, include_content: bool = False, case_sensitive: bool = False, use_regex: bool = False, file_extensions: Optional[List[str]] = None, max_scan_entries: int = 120000, max_file_size_mb: int = 5) -> Dict[str, Any]:
        clean_query = str(query or "").strip()
        if not clean_query: return _fail("Deep search query was empty.", "empty_query")

        search_roots = self._resolve_search_roots(roots=roots, include_all_drives=include_all_drives)
        if not search_roots: return _fail("No accessible roots were found for deep search.", "missing_roots")

        max_hits = max(1, int(max_results))
        max_scanned = max(1, int(max_scan_entries))
        max_size_bytes = max(1, int(max_file_size_mb)) * 1024 * 1024

        normalized_extensions = set()
        for ext in file_extensions or []:
            clean_ext = str(ext or "").strip().lower()
            if not clean_ext: continue
            if not clean_ext.startswith("."): clean_ext = "." + clean_ext
            normalized_extensions.add(clean_ext)

        regex = None
        if use_regex:
            flags = 0 if case_sensitive else re.IGNORECASE
            try: regex = re.compile(clean_query, flags=flags)
            except re.error as exc: return _fail("Deep search regex is invalid.", str(exc))

        needle = clean_query if case_sensitive else clean_query.lower()

        def _matches(text: str) -> bool:
            if regex is not None: return bool(regex.search(text))
            if case_sensitive: return needle in text
            return needle in text.lower()

        def _snippet(text: str) -> str:
            if not text: return ""
            if regex is not None:
                match = regex.search(text)
                if not match: return ""
                start_idx, end_idx = match.start(), match.end()
            else:
                haystack = text if case_sensitive else text.lower()
                start_idx = haystack.find(needle)
                if start_idx < 0: return ""
                end_idx = start_idx + len(needle)

            left = max(0, start_idx - 80)
            right = min(len(text), end_idx + 140)
            return re.sub(r"\s+", " ", text[left:right]).strip()

        scanned_entries = 0
        truncated = False
        results: List[Dict[str, Any]] = []

        skip_dir_names = {"$recycle.bin", "system volume information", ".git", ".venv", "node_modules", "__pycache__"}

        for root in search_roots:
            if len(results) >= max_hits or scanned_entries >= max_scanned:
                truncated = True
                break

            for current_root, dirs, files in os.walk(root, topdown=True, onerror=lambda _e: None):
                if len(results) >= max_hits or scanned_entries >= max_scanned:
                    truncated = True
                    break
                current_path = Path(current_root)

                kept_dirs = []
                for dirname in dirs:
                    if dirname.lower() in skip_dir_names: continue
                    scanned_entries += 1
                    if scanned_entries > max_scanned:
                        truncated = True
                        break

                    dir_path = current_path / dirname
                    kept_dirs.append(dirname)
                    if _matches(str(dir_path)):
                        results.append({"path": str(dir_path.resolve()), "type": "directory", "match": "name"})
                        if len(results) >= max_hits:
                            truncated = True
                            break

                dirs[:] = kept_dirs
                if truncated and (len(results) >= max_hits or scanned_entries >= max_scanned): break

                for filename in files:
                    scanned_entries += 1
                    if scanned_entries > max_scanned:
                        truncated = True
                        break

                    file_path = current_path / filename
                    if _matches(str(file_path)):
                        results.append({"path": str(file_path.resolve()), "type": "file", "match": "name"})
                        if len(results) >= max_hits:
                            truncated = True
                            break

                    if not include_content: continue

                    suffix = file_path.suffix.lower()
                    if normalized_extensions and suffix not in normalized_extensions: continue
                    if not normalized_extensions and suffix and suffix not in TEXT_SEARCH_EXTENSIONS: continue

                    try:
                        if file_path.stat().st_size > max_size_bytes: continue
                    except Exception: continue

                    try:
                        file_text = file_path.read_text(encoding="utf-8", errors="ignore")
                    except Exception: continue

                    if _matches(file_text):
                        entry: Dict[str, Any] = {"path": str(file_path.resolve()), "type": "file", "match": "content"}
                        context = _snippet(file_text)
                        if context: entry["snippet"] = context
                        results.append(entry)
                        if len(results) >= max_hits:
                            truncated = True
                            break

                if truncated and (len(results) >= max_hits or scanned_entries >= max_scanned): break

        return _ok(f"Deep search found {len(results)} match(es).", data={"query": clean_query, "results": results, "scanned_entries": scanned_entries, "truncated": truncated, "roots": [str(root) for root in search_roots]})

    def analyze_path(self, path: str, max_items: int = 3500) -> Dict[str, Any]:
        clean_path = str(path or "").strip()
        if not clean_path: return _fail("Analysis path was empty.", "empty_path")

        try:
            target = _safe_resolve_path(clean_path)
            if not target.exists(): return _fail("Path was not found.", f"missing_path: {target}")

            if target.is_file():
                stat = target.stat()
                return _ok(f"Analyzed file: {target.name}", data={"path": str(target), "type": "file", "size_bytes": stat.st_size, "suffix": target.suffix.lower(), "modified_epoch": stat.st_mtime})

            cap = max(200, int(max_items))
            scanned_files, total_dirs, total_size = 0, 0, 0
            truncated = False

            extension_counts: Dict[str, int] = {}
            largest_files: List[Dict[str, Any]] = []

            for current_root, dirs, files in os.walk(target, topdown=True, onerror=lambda _e: None):
                total_dirs += len(dirs)
                for filename in files:
                    scanned_files += 1
                    if scanned_files > cap:
                        truncated = True
                        break

                    file_path = Path(current_root) / filename
                    try: size = file_path.stat().st_size
                    except Exception: continue

                    total_size += size
                    suffix = file_path.suffix.lower() or "<none>"
                    extension_counts[suffix] = extension_counts.get(suffix, 0) + 1

                    largest_files.append({"path": str(file_path.resolve()), "size_bytes": size})
                    largest_files.sort(key=lambda item: int(item["size_bytes"]), reverse=True)
                    if len(largest_files) > 8: largest_files = largest_files[:8]
                if truncated: break

            top_extensions = sorted(extension_counts.items(), key=lambda item: item[1], reverse=True)[:12]
            return _ok(f"Analyzed directory: {target}", data={"path": str(target), "type": "directory", "total_files": scanned_files if not truncated else cap, "total_dirs": total_dirs, "total_size_bytes": total_size, "top_extensions": top_extensions, "largest_files": largest_files, "truncated": truncated, "max_items": cap})
        except Exception as exc:
            return _fail("Path analysis failed.", str(exc))

    def create_path(self, path: str, kind: str = "file", content: str = "", overwrite: bool = False) -> Dict[str, Any]:
        clean_path = str(path or "").strip()
        if not clean_path: return _fail("Create path was empty.", "empty_path")
        clean_kind = str(kind or "file").strip().lower()

        try:
            target = _safe_resolve_path(clean_path)
            if _is_protected_path(target): return _fail("Path is protected by safe mode.", "protected_path")
            if clean_kind in {"dir", "folder", "directory"}:
                if target.exists() and not target.is_dir(): return _fail("Target exists as a file.", f"path_conflict: {target}")
                target.mkdir(parents=True, exist_ok=True)
                return _ok("Directory is ready.", data={"path": str(target), "type": "directory"})

            target.parent.mkdir(parents=True, exist_ok=True)
            if target.exists() and not overwrite:
                return _fail("File already exists. Use overwrite=true to replace it.", f"existing_file: {target}")

            target.write_text(str(content or ""), encoding="utf-8")
            return _ok("File created successfully.", data={"path": str(target), "type": "file"})
        except Exception as exc:
            return _fail("Create path failed.", str(exc))

    def move_path(self, src: str, dst: str, overwrite: bool = False) -> Dict[str, Any]:
        clean_src = str(src or "").strip()
        clean_dst = str(dst or "").strip()
        if not clean_src or not clean_dst: return _fail("Move source or destination was empty.", "empty_move_paths")

        try:
            src_path = _safe_resolve_path(clean_src)
            dst_path = _safe_resolve_path(clean_dst)
            if _is_protected_path(src_path) or _is_protected_path(dst_path): return _fail("Path is protected by safe mode.", "protected_path")
            if not src_path.exists(): return _fail("Move source was not found.", f"missing_source: {src_path}")

            if dst_path.exists():
                if not overwrite: return _fail("Move destination already exists. Use overwrite=true to replace it.", f"existing_destination: {dst_path}")
                if dst_path.is_dir(): shutil.rmtree(dst_path)
                else: dst_path.unlink()

            dst_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(str(src_path), str(dst_path))
            return _ok("Move completed.", data={"from": str(src_path), "to": str(dst_path)})
        except Exception as exc:
            return _fail("Move operation failed.", str(exc))

    def copy_path(self, src: str, dst: str, overwrite: bool = False) -> Dict[str, Any]:
        clean_src = str(src or "").strip()
        clean_dst = str(dst or "").strip()
        if not clean_src or not clean_dst: return _fail("Copy source or destination was empty.", "empty_copy_paths")

        try:
            src_path = _safe_resolve_path(clean_src)
            dst_path = _safe_resolve_path(clean_dst)
            if _is_protected_path(src_path) or _is_protected_path(dst_path): return _fail("Path is protected by safe mode.", "protected_path")
            if not src_path.exists(): return _fail("Copy source was not found.", f"missing_source: {src_path}")

            if dst_path.exists():
                if not overwrite: return _fail("Copy destination already exists. Use overwrite=true to replace it.", f"existing_destination: {dst_path}")
                if dst_path.is_dir(): shutil.rmtree(dst_path)
                else: dst_path.unlink()

            dst_path.parent.mkdir(parents=True, exist_ok=True)
            if src_path.is_dir(): shutil.copytree(src_path, dst_path)
            else: shutil.copy2(src_path, dst_path)
            return _ok("Copy completed.", data={"from": str(src_path), "to": str(dst_path)})
        except Exception as exc:
            return _fail("Copy operation failed.", str(exc))

    def delete_path(self, path: str, recursive: bool = False, use_trash: bool = True) -> Dict[str, Any]:
        clean_path = str(path or "").strip()
        if not clean_path: return _fail("Delete path was empty.", "empty_path")

        try:
            target = _safe_resolve_path(clean_path)
            if _is_protected_path(target): return _fail("Path is protected by safe mode.", "protected_path")
            if not target.exists(): return _fail("Delete target was not found.", f"missing_path: {target}")

            if _to_bool(use_trash, default=True) and send2trash is not None:
                send2trash(str(target))
                return _ok("Moved target to recycle bin.", data={"path": str(target)})

            if target.is_dir():
                if not _to_bool(recursive, default=False): return _fail("Target is a directory. Set recursive=true to delete it permanently.", "directory_requires_recursive")
                shutil.rmtree(target)
            else:
                target.unlink()
            return _ok("Delete completed.", data={"path": str(target)})
        except Exception as exc:
            return _fail("Delete operation failed.", str(exc))

    @staticmethod
    def _normalize_app_alias(name: str) -> str:
        clean = str(name or "").strip().lower()
        clean = re.sub(r"[^a-z0-9\s\-_.]+", " ", clean)
        clean = re.sub(r"\b(app|application|software|program)\b", " ", clean)
        return re.sub(r"\s+", " ", clean).strip()   

    def launch_application(self, app_name: str, args: Optional[List[str]] = None, target_path: str = "") -> Dict[str, Any]:
        clean_app = str(app_name or "").strip()
        clean_target = str(target_path or "").strip()
        clean_args = [str(item).strip() for item in (args or []) if str(item).strip()]

        normalized_app = self._normalize_app_alias(clean_app)
        
        # 1. Windows Aliases and specific application paths
        if os.name == "nt" and normalized_app in WINDOWS_APP_ALIASES:
            executable = WINDOWS_APP_ALIASES[normalized_app]
            if ":\\" in executable:
                if Path(executable).exists():
                    try:
                        subprocess.Popen([executable] + clean_args, creationflags=0x00000008 | 0x00000200)
                        return _ok(f"Launched exact path alias: {executable}")
                    except Exception: pass
            else:
                try:
                    args_str = " ".join(f'"{a}"' for a in clean_args)
                    target_str = f'"{clean_target}"' if clean_target else ""
                    cmd = f'cmd /c start "" "{executable}" {args_str} {target_str}'.strip()
                    subprocess.Popen(cmd, shell=True)
                    return _ok(f"Launched known alias: {executable}")
                except Exception: pass 
        
        # 2. AppOpener integration
        if os.name == "nt" and APPOPENER_AVAILABLE and appopener_open is not None:
            try:
                appopener_open(clean_app, match_closest=True, output=False, throw_error=True)
                return _ok(f"Launched via AppOpener: {clean_app}")
            except Exception: pass    

        # 3. Web Service fallbacks
        web_url = WEB_APP_ALIASES.get(normalized_app)
        if web_url and not clean_target:
            try:
                if os.name == "nt": os.startfile(web_url)
                else: webbrowser.open(web_url)
                return _ok(f"Opened {clean_app} in browser.", data={"url": web_url})
            except Exception as exc: return _fail("Launch web application failed.", str(exc))

        # 4. Direct Executable Path 
        app_path = Path(clean_app).expanduser()
        if app_path.exists():
            command = [str(app_path.resolve())] + clean_args
            if clean_target: command.append(str(_safe_resolve_path(clean_target)))
            try:
                if os.name == "nt": subprocess.Popen(command, creationflags=0x00000008 | 0x00000200)
                else: subprocess.Popen(command)
                return _ok(f"Launched specific path: {clean_app}")
            except Exception: pass

        # 5. Native OS Fallback
        try:
            executable = clean_app
            if os.name == "nt":
                resolved_exec = shutil.which(executable)
                creation_flags = 0x00000008 | 0x00000200
                if resolved_exec:
                    command = [resolved_exec] + clean_args
                    if clean_target: command.append(str(_safe_resolve_path(clean_target)))
                    subprocess.Popen(command, creationflags=creation_flags)
                else:
                    target_str = f'"{clean_target}"' if clean_target else ""
                    subprocess.Popen(f'cmd /c start "" "{executable}" {target_str}'.strip(), shell=True, creationflags=creation_flags)
            else:
                command = ["open", "-a", executable] + clean_args if SYSTEM_NAME == "darwin" else [executable] + clean_args
                if clean_target: command.append(str(_safe_resolve_path(clean_target)))
                subprocess.Popen(command)

            return _ok(f"Launched application fallback: {clean_app}")
        except Exception as exc:
            return _fail("Launch application completely failed.", str(exc))

    def close_application(self, app_name: str, force: bool = True) -> Dict[str, Any]:
        clean_app = str(app_name or "").strip()
        if not clean_app: return _fail("Application name was empty.", "empty_app_name")

        try:
            target = clean_app
            if os.name == "nt":
                normalized_app = self._normalize_app_alias(clean_app)
                alias = WINDOWS_APP_ALIASES.get(normalized_app)
                if alias and ":\\" not in alias: 
                    exe_name = alias if alias.lower().endswith(".exe") else alias + ".exe"
                    command = ["taskkill", "/IM", exe_name]
                    if _to_bool(force, default=True): command.append("/F")
                    completed = subprocess.run(command, capture_output=True, text=True, errors="replace", check=False)
                    if completed.returncode == 0: return _ok(f"Closed known alias: {exe_name}")

                if APPOPENER_AVAILABLE and appopener_close is not None:
                    try:
                        appopener_close(clean_app, match_closest=True, output=False, throw_error=True)
                        return _ok(f"Close command sent via AppOpener: {clean_app}")
                    except Exception: pass

                exe_name = target if target.lower().endswith(".exe") else target + ".exe"
                command = ["taskkill", "/IM", exe_name]
                if _to_bool(force, default=True): command.append("/F")
                completed = subprocess.run(command, capture_output=True, text=True, errors="replace", check=False)
                output = ((completed.stdout or "") + "\n" + (completed.stderr or "")).strip()
                if completed.returncode == 0: return _ok(f"Closed application: {clean_app}")
                return _fail("Close application failed.", output[:1200] or f"exit_code={completed.returncode}")

            command = ["pkill", "-f", target]
            completed = subprocess.run(command, capture_output=True, text=True, check=False)
            if completed.returncode in {0, 1}: return _ok(f"Close request sent for application: {clean_app}")
            output = ((completed.stdout or "") + "\n" + (completed.stderr or "")).strip()
            return _fail("Close application failed.", output[:1200] or f"exit_code={completed.returncode}")
        except Exception as exc:
            return _fail("Close application failed.", str(exc))

    def list_running_applications(self, max_results: int = 120) -> Dict[str, Any]:
        cap = max(1, int(max_results))
        try:
            apps: List[str] = []
            if os.name == "nt":
                completed = subprocess.run(["tasklist", "/FO", "CSV", "/NH"], capture_output=True, text=True, errors="replace", check=False)
                if completed.returncode != 0: return _fail("Could not list running applications.", (completed.stderr or "").strip())
                reader = csv.reader((completed.stdout or "").splitlines())
                for row in reader:
                    if not row: continue
                    name = str(row[0]).strip()
                    if name: apps.append(name)
                    if len(apps) >= cap: break
            else:
                completed = subprocess.run(["ps", "-e", "-o", "comm="], capture_output=True, text=True, check=False)
                if completed.returncode != 0: return _fail("Could not list running applications.", (completed.stderr or "").strip())
                for line in (completed.stdout or "").splitlines():
                    name = line.strip()
                    if name: apps.append(name)
                    if len(apps) >= cap: break
            return _ok(f"Found {len(apps)} running application(s).", data=apps)
        except Exception as exc:
            return _fail("Listing running applications failed.", str(exc))

    def open_service(self, service: str) -> Dict[str, Any]:
        clean = self._normalize_app_alias(service)
        if not clean: return _fail("Service name was empty.", "empty_service")

        service_urls = {
            "gmail": "https://mail.google.com",
            "outlook": "https://outlook.live.com/mail/",
            "whatsapp": "https://web.whatsapp.com",
            "telegram": "https://web.telegram.org",
        }

        if clean in service_urls:
            try:
                url = service_urls[clean]
                webbrowser.open(url)
                return _ok(f"Opened {clean}.", data={"url": url})
            except Exception as exc: return _fail(f"Failed to open {clean}.", str(exc))
        return self.launch_application(clean)

    def read_file_text(self, file_path: str, max_chars: int = 15000) -> Dict[str, Any]:
        try:
            path = Path(file_path).expanduser().resolve()
            if not path.exists() or not path.is_file(): return _fail("File was not found.", f"missing_file: {path}")

            suffix = path.suffix.lower()
            if suffix == ".pdf":
                if PyPDF2 is None: return _fail("PyPDF2 is not available to read PDF files.", "missing_dependency: PyPDF2")
                try:
                    with path.open("rb") as fh:
                        reader = PyPDF2.PdfReader(fh)
                        pages_text = []
                        for page in reader.pages:
                            page_text = page.extract_text() or ""
                            pages_text.append(page_text)
                            if sum(len(t) for t in pages_text) >= max_chars: break
                        text = "\n".join(pages_text)
                except Exception as exc: return _fail("Could not parse PDF file.", str(exc))
                    
            elif suffix in {".doc", ".docx"}:
                if docx is None: return _fail("python-docx is not available to read DOCX files.", "missing_dependency: python-docx")
                try:
                    document = Document(str(path))
                    text = "\n".join(p.text for p in document.paragraphs)
                except Exception as exc: return _fail("Could not parse DOCX file.", str(exc))
            else:
                try: text = path.read_text(encoding="utf-8", errors="ignore")
                except Exception: text = path.read_bytes().decode("utf-8", errors="ignore")

            text = text[:max_chars]
            return _ok(f"Read file successfully: {path.name}", data={"path": str(path), "text": text})
        except Exception as exc:
            return _fail("File read failed.", str(exc))

    def open_file(self, file_path: str, resolve_by_hint: bool = False, roots: Optional[List[str]] = None, include_content: bool = True) -> Dict[str, Any]:
        try:
            raw_target = str(file_path or "").strip()
            if not raw_target: return _fail("Cannot open path because it was empty.", "empty_path")

            path = Path(raw_target).expanduser().resolve()
            resolved_from_hint = False

            if not path.exists() and _to_bool(resolve_by_hint, default=False):
                search_result = self.smart_heuristic_search_files(hint=raw_target, roots=roots, max_results=1, include_content=_to_bool(include_content, default=True))
                candidates = search_result.get("data") if isinstance(search_result, dict) else []
                if candidates:
                    path = Path(str(candidates[0])).expanduser().resolve()
                    resolved_from_hint = True

            if not path.exists(): return _fail("Cannot open path because it does not exist.", f"missing_file: {path}")

            if os.name == "nt": os.startfile(str(path))
            elif os.name == "posix": subprocess.Popen(["xdg-open", str(path)])
            else: subprocess.Popen(["open", str(path)])

            if resolved_from_hint: return _ok(f"Opened path: {path.name} (resolved from hint).", data=str(path))
            return _ok(f"Opened path: {path.name}", data=str(path))
        except Exception as exc:
            return _fail("Failed to open path.", str(exc))

    def move_mouse(self, x: int, y: int, duration: float = 0.2) -> Dict[str, Any]:
        try:
            if pyautogui is None: return _fail("PyAutoGUI is unavailable.", "missing_dependency: pyautogui")
            pyautogui.moveTo(int(x), int(y), duration=max(0.0, duration))
            return _ok(f"Mouse moved to ({x}, {y}).")
        except Exception as exc: return _fail("Mouse move failed.", str(exc))

    def click(self, button: str = "left", clicks: int = 1) -> Dict[str, Any]:
        try:
            if pyautogui is None: return _fail("PyAutoGUI is unavailable.", "missing_dependency: pyautogui")
            pyautogui.click(button=button, clicks=max(1, int(clicks)))
            return _ok(f"Mouse click completed ({button}, {clicks}x).")
        except Exception as exc: return _fail("Mouse click failed.", str(exc))

    def type_text(self, text: str, interval: float = 0.01) -> Dict[str, Any]:
        try:
            if pyautogui is None: return _fail("PyAutoGUI is unavailable.", "missing_dependency: pyautogui")
            pyautogui.write(str(text), interval=max(0.0, float(interval)))
            return _ok("Text typing completed.")
        except Exception as exc: return _fail("Typing action failed.", str(exc))

    def press_key(self, key: str) -> Dict[str, Any]:
        try:
            if pyautogui is None: return _fail("PyAutoGUI is unavailable.", "missing_dependency: pyautogui")
            pyautogui.press(str(key))
            return _ok(f"Pressed key: {key}")
        except Exception as exc: return _fail("Key press failed.", str(exc))

    def hotkey(self, keys: List[str]) -> Dict[str, Any]:
        try:
            if pyautogui is None: return _fail("PyAutoGUI is unavailable.", "missing_dependency: pyautogui")
            if not keys: return _fail("No hotkey keys provided.", "empty_hotkey")
            pyautogui.hotkey(*[str(k) for k in keys])
            return _ok(f"Hotkey sent: {' + '.join(keys)}")
        except Exception as exc: return _fail("Hotkey action failed.", str(exc))

    def run_system_command(self, command: str, timeout: int = 25) -> Dict[str, Any]:
        clean = (command or "").strip()
        if not clean: return _fail("Command text was empty.", "empty_command")
        try:
            env = os.environ.copy()
            env.setdefault("OLLAMA_FLASH_ATTENTION", "1")
            env.setdefault("OLLAMA_KV_CACHE_TYPE", "q4_0")
            completed = subprocess.run(clean, shell=True, capture_output=True, text=True, timeout=max(1, int(timeout)), env=env)
            combined = (completed.stdout or "") + ("\n" + completed.stderr if completed.stderr else "")
            compact = combined.strip()[:4000]
            if completed.returncode == 0: return _ok("System command completed.", data={"output": compact})
            return _fail("System command failed.", f"exit_code={completed.returncode}; output={compact}")
        except Exception as exc: return _fail("System command execution failed.", str(exc))

    def run_open_interpreter(self, instruction: str, timeout: int = 120) -> Dict[str, Any]:
        clean = (instruction or "").strip()
        if not clean: return _fail("Open Interpreter instruction was empty.", "empty_instruction")

        exe = shutil.which("interpreter")
        if not exe: return _fail("Open Interpreter CLI not found.", "missing_dependency: open-interpreter")

        try:
            completed = subprocess.run([exe, "--auto-run"], input=clean, capture_output=True, text=True, timeout=max(5, int(timeout)))
            combined = (completed.stdout or "") + ("\n" + completed.stderr if completed.stderr else "")
            compact = combined.strip()[:4000]
            if completed.returncode == 0: return _ok("Open Interpreter completed.", data={"output": compact})
            return _fail("Open Interpreter failed.", f"exit_code={completed.returncode}; output={compact}")
        except Exception as exc: return _fail("Open Interpreter execution failed.", str(exc))

    def toggle_dark_mode(self, enable: Optional[bool] = None) -> Dict[str, Any]:
        if os.name == "nt": return self._toggle_windows_dark_mode(enable)
        return self._toggle_linux_dark_mode(enable)

    def _toggle_windows_dark_mode(self, enable: Optional[bool]) -> Dict[str, Any]:
        key = r"HKCU\Software\Microsoft\Windows\CurrentVersion\Themes\Personalize"
        current_is_dark = False
        try:
            probe = subprocess.run(["reg", "query", key, "/v", "AppsUseLightTheme"], capture_output=True, text=True)
            current_is_dark = "0x0" in (probe.stdout or "")
        except Exception: current_is_dark = False

        target_is_dark = (not current_is_dark) if enable is None else bool(enable)
        value = "0" if target_is_dark else "1"

        try:
            subprocess.run(["reg", "add", key, "/v", "AppsUseLightTheme", "/t", "REG_DWORD", "/d", value, "/f"], check=False, capture_output=True, text=True)
            subprocess.run(["reg", "add", key, "/v", "SystemUsesLightTheme", "/t", "REG_DWORD", "/d", value, "/f"], check=False, capture_output=True, text=True)
            mode = "dark" if target_is_dark else "light"
            return _ok(f"Windows theme toggled to {mode} mode.")
        except Exception as exc: return _fail("Failed to toggle Windows dark mode.", str(exc))

    def _toggle_linux_dark_mode(self, enable: Optional[bool]) -> Dict[str, Any]:
        target_is_dark = True if enable is None else bool(enable)
        color_scheme = "prefer-dark" if target_is_dark else "default"
        gtk_theme = "Adwaita-dark" if target_is_dark else "Adwaita"

        try:
            subprocess.run(["gsettings", "set", "org.gnome.desktop.interface", "color-scheme", color_scheme], check=False, capture_output=True, text=True)
            subprocess.run(["gsettings", "set", "org.gnome.desktop.interface", "gtk-theme", gtk_theme], check=False, capture_output=True, text=True)
            mode = "dark" if target_is_dark else "light"
            return _ok(f"Linux theme toggled to {mode} mode.")
        except Exception as exc: return _fail("Failed to toggle Linux dark mode.", str(exc))

    def _resolve_attachment_paths(self, attachments: Optional[List[str]], roots: Optional[List[str]] = None) -> tuple[List[str], List[str]]:
        resolved: List[str] = []
        unresolved: List[str] = []

        for item in attachments or []:
            hint = str(item or "").strip()
            if not hint: continue
            try:
                direct_path = _safe_resolve_path(hint)
                if direct_path.exists() and direct_path.is_file():
                    resolved.append(str(direct_path))
                    continue
            except Exception: pass

            found = self.search_files_by_hint(hint, roots=roots, max_results=1)
            candidates = found.get("data") if isinstance(found, dict) else []
            if candidates: resolved.append(str(candidates[0]))
            else: unresolved.append(hint)

        unique_resolved: List[str] = []
        seen = set()
        for path in resolved:
            if path not in seen:
                unique_resolved.append(path)
                seen.add(path)
        return unique_resolved, unresolved

    @staticmethod
    def _smtp_default(provider: str) -> tuple[str, int]:
        clean = str(provider or "").strip().lower()
        if clean == "gmail": return "smtp.gmail.com", 587
        if clean in {"outlook", "hotmail", "live"}: return "smtp.office365.com", 587
        return "", 587

    def send_email(self, to_email: str, subject: str, body: str, provider: str = "gmail", attachments: Optional[List[str]] = None, roots: Optional[List[str]] = None, smtp_host: str = "", smtp_port: int = 587, smtp_user: str = "", smtp_password: str = "", from_email: str = "", send_now: bool = False, use_outlook_desktop: bool = True) -> Dict[str, Any]:
        clean_to = str(to_email or "").strip()
        if not clean_to: return _fail("Recipient email is required.", "missing_to_email")

        clean_provider = str(provider or "gmail").strip().lower() or "gmail"
        resolved_attachments, unresolved_attachments = self._resolve_attachment_paths(attachments, roots=roots)

        if clean_provider == "outlook" and os.name == "nt" and win32 is not None and _to_bool(use_outlook_desktop, default=True):
            try:
                def run_outlook_async():
                    import pythoncom
                    # Crucial: Initialize COM library lifecycle for this background thread
                    pythoncom.CoInitialize() 
                    try:
                        outlook = win32.Dispatch("Outlook.Application")
                        mail = outlook.CreateItem(0)
                        mail.To = clean_to
                        mail.Subject = str(subject or "").strip()
                        mail.Body = str(body or "").strip()
                        for file_path in resolved_attachments: 
                            mail.Attachments.Add(file_path)

                        if _to_bool(send_now, default=False):
                            mail.Send()
                        else:
                            mail.Save()
                    except Exception as thread_exc:
                        print(f"[Async Outlook Error] Automation failed: {thread_exc}")
                    finally:
                        pythoncom.CoUninitialize()

                # Offload to background thread immediately
                threading.Thread(target=run_outlook_async, daemon=True).start()
                
                status = "Outlook automation triggered in the background."
                payload = {"provider": "outlook_desktop_async", "to": clean_to, "attachments": resolved_attachments}
                if unresolved_attachments: 
                    payload["unresolved_attachments"] = unresolved_attachments
                return _ok(status, data=payload)
                
            except Exception as exc: 
                return _fail("Outlook desktop thread initialization failed.", str(exc))

        default_host, default_port = self._smtp_default(clean_provider)
        final_host = str(smtp_host or default_host).strip()
        final_port = int(smtp_port or default_port or 587)
        sender = str(from_email or smtp_user or "").strip()

        if not final_host: return _fail("SMTP host is required for this email provider.", "missing_smtp_host")
        if not sender: return _fail("Sender address is required (from_email or smtp_user).", "missing_sender")

        try:
            msg = EmailMessage()
            msg["From"] = sender
            msg["To"] = clean_to
            msg["Subject"] = str(subject or "").strip()
            msg.set_content(str(body or "").strip())

            for file_path in resolved_attachments:
                source = Path(file_path)
                mime_type, _ = mimetypes.guess_type(str(source))
                if not mime_type: mime_type = "application/octet-stream"
                maintype, subtype = mime_type.split("/", 1)
                msg.add_attachment(source.read_bytes(), maintype=maintype, subtype=subtype, filename=source.name)

            if _to_bool(send_now, default=False):
                if not str(smtp_user or "").strip() or not str(smtp_password or ""):
                    return _fail("SMTP credentials are required to send now.", "missing_smtp_credentials")
                with smtplib.SMTP(final_host, final_port, timeout=25) as smtp:
                    smtp.starttls()
                    smtp.login(str(smtp_user).strip(), str(smtp_password))
                    smtp.send_message(msg)

                payload = {"provider": clean_provider, "mode": "smtp_send", "smtp_host": final_host, "to": clean_to, "attachments": resolved_attachments}
                if unresolved_attachments: payload["unresolved_attachments"] = unresolved_attachments
                return _ok("Email sent via SMTP.", data=payload)

            draft_dir = Path.cwd() / "cache" / "mail_drafts"
            draft_dir.mkdir(parents=True, exist_ok=True)
            draft_name = f"email_draft_{int(time.time())}.eml"
            draft_path = draft_dir / draft_name
            draft_path.write_bytes(bytes(msg))

            payload = {"provider": clean_provider, "mode": "draft_file", "draft": str(draft_path.resolve()), "to": clean_to, "attachments": resolved_attachments}
            if unresolved_attachments: payload["unresolved_attachments"] = unresolved_attachments
            return _ok("Email draft file created.", data=payload)
        except Exception as exc: return _fail("Email operation failed.", str(exc))

    def draft_email_attachment(self, to_email: str, subject: str, body: str, file_hint: str, roots: Optional[List[str]] = None, smtp_host: str = "", smtp_port: int = 587, smtp_user: str = "", smtp_password: str = "", from_email: str = "", send_now: bool = False) -> Dict[str, Any]:
        return self.send_email(to_email=to_email, subject=subject, body=body, provider="outlook" if os.name == "nt" else "custom", attachments=[str(file_hint or "")], roots=roots, smtp_host=smtp_host, smtp_port=smtp_port, smtp_user=smtp_user, smtp_password=smtp_password, from_email=from_email, send_now=send_now, use_outlook_desktop=True)

    def send_telegram(self, bot_token: str, chat_id: str, message: str = "", file_path: str = "", file_hint: str = "", roots: Optional[List[str]] = None, disable_web_preview: bool = False) -> Dict[str, Any]:
        if requests is None: return _fail("requests is unavailable for Telegram messaging.", "missing_dependency: requests")

        clean_token = str(bot_token or "").strip()
        clean_chat = str(chat_id or "").strip()
        clean_message = str(message or "").strip()
        if not clean_token or not clean_chat: return _fail("Telegram bot token and chat_id are required.", "missing_telegram_credentials")

        base_url = f"https://api.telegram.org/bot{clean_token}"
        sent = []

        try:
            if clean_message:
                payload = {"chat_id": clean_chat, "text": clean_message, "disable_web_page_preview": _to_bool(disable_web_preview, default=False)}
                response = requests.post(f"{base_url}/sendMessage", data=payload, timeout=20)
                if not response.ok: return _fail("Telegram message send failed.", response.text[:1200])
                sent.append("message")

            file_candidate = str(file_path or file_hint or "").strip()
            if file_candidate:
                resolved, unresolved = self._resolve_attachment_paths([file_candidate], roots=roots)
                if not resolved: return _fail("Telegram file was not found.", f"unresolved_file: {unresolved}")
                source = Path(resolved[0])
                with source.open("rb") as handle:
                    payload = {"chat_id": clean_chat}
                    if clean_message: payload["caption"] = clean_message[:900]
                    response = requests.post(f"{base_url}/sendDocument", data=payload, files={"document": handle}, timeout=40)
                if not response.ok: return _fail("Telegram file send failed.", response.text[:1200])
                sent.append("document")

            if not sent: return _fail("Nothing to send to Telegram. Provide message or file_path.", "empty_telegram_payload")
            return _ok("Telegram send completed.", data={"chat_id": clean_chat, "sent": sent})
        except Exception as exc: return _fail("Telegram operation failed.", str(exc))

    def send_whatsapp(self, to_number: str, message: str, use_twilio: bool = False, twilio_account_sid: str = "", twilio_auth_token: str = "", twilio_from: str = "", media_url: str = "") -> Dict[str, Any]:
        clean_to = str(to_number or "").strip()
        clean_message = str(message or "").strip()
        clean_media_url = str(media_url or "").strip()

        if not clean_to: return _fail("WhatsApp recipient number is required.", "missing_whatsapp_number")
        wants_twilio = _to_bool(use_twilio, default=False) or bool(str(twilio_account_sid or "").strip() and str(twilio_auth_token or "").strip() and str(twilio_from or "").strip())

        if wants_twilio:
            if requests is None: return _fail("requests is unavailable for Twilio WhatsApp.", "missing_dependency: requests")
            sid = str(twilio_account_sid or "").strip()
            token = str(twilio_auth_token or "").strip()
            from_number = str(twilio_from or "").strip()
            if not sid or not token or not from_number: return _fail("Twilio SID, token, and from number are required.", "missing_twilio_credentials")

            to_value = clean_to if clean_to.lower().startswith("whatsapp:") else f"whatsapp:{clean_to}"
            from_value = from_number if from_number.lower().startswith("whatsapp:") else f"whatsapp:{from_number}"
            payload = {"To": to_value, "From": from_value}
            if clean_message: payload["Body"] = clean_message
            if clean_media_url: payload["MediaUrl"] = clean_media_url

            try:
                response = requests.post(f"https://api.twilio.com/2010-04-01/Accounts/{sid}/Messages.json", data=payload, auth=(sid, token), timeout=25)
                if response.status_code >= 400: return _fail("Twilio WhatsApp send failed.", response.text[:1200])
                response_data = response.json() if hasattr(response, "json") else {}
                return _ok("WhatsApp message sent via Twilio.", data={"to": to_value, "sid": response_data.get("sid", "")})
            except Exception as exc: return _fail("Twilio WhatsApp send failed.", str(exc))

        if not clean_message: return _fail("WhatsApp web mode supports text messages only. Provide a message.", "missing_message")

        if pywhatkit is not None:
            try:
                pywhatkit.sendwhatmsg_instantly(phone_no=clean_to, message=clean_message, wait_time=15, tab_close=True, close_time=4)
                return _ok("WhatsApp Web message queued. Ensure your browser is logged in to WhatsApp.", data={"to": clean_to})
            except Exception as exc: return _fail("WhatsApp Web send failed.", str(exc))

        try:
            url = f"https://web.whatsapp.com/send?phone={quote_plus(clean_to)}&text={quote_plus(clean_message)}"
            webbrowser.open(url)
            return _ok("Opened WhatsApp Web compose window.", data={"url": url})
        except Exception as exc: return _fail("Failed to open WhatsApp Web.", str(exc))

    def execute_action(self, action: Dict[str, Any]) -> Dict[str, Any]:
        try:
            action_name = str(action.get("action", "")).strip().lower()
            roots = _to_str_list(action.get("roots")) or None
            
            # --- Application Handlers ---
            if action_name in {"launch_application", "open"}:
                target_app = str(action.get("app") or action.get("name") or action.get("target") or "")
                return self.launch_application(app_name=target_app, args=_to_str_list(action.get("args")) or None, target_path=str(action.get("path") or action.get("target_path") or ""))
            if action_name in {"close_application", "close"}:
                target_app = str(action.get("app") or action.get("name") or action.get("target") or "")
                return self.close_application(app_name=target_app, force=_to_bool(action.get("force", True), default=True))
            if action_name in {"list_running_apps", "list_running_applications"}:
                return self.list_running_applications(max_results=int(action.get("max_results", 120) or 120))
            if action_name == "open_service":
                return self.open_service(str(action.get("service", "")))

            # --- File System Operations ---
            if action_name == "list_system_roots":
                return self.list_system_roots()
            if action_name == "search_mirror":
                return self.search_mirror(query=str(action.get("query") or action.get("hint") or ""), limit=int(action.get("max_results", 8) or 8))
            if action_name == "list_directory":
                mirror_context = self._get_mirror_context(self._mirror_query_from_action(action, action.get("path") or ""), limit=int(action.get("mirror_limit", 6) or 6)) if self._mirror_requested(action) else None
                result = self.list_directory(path=str(action.get("path", "")), recursive=_to_bool(action.get("recursive", False), default=False), max_depth=int(action.get("max_depth", 2) or 2), max_results=int(action.get("max_results", 300) or 300), pattern=str(action.get("pattern", "")), include_files=_to_bool(action.get("include_files", True), default=True), include_dirs=_to_bool(action.get("include_dirs", True), default=True))
                return self._attach_mirror_context(result, mirror_context)
            if action_name in {"deep_search", "deep_search_paths"}:
                mirror_context = self._get_mirror_context(self._mirror_query_from_action(action, action.get("query") or action.get("hint") or ""), limit=int(action.get("mirror_limit", 6) or 6)) if self._mirror_requested(action) else None
                result = self.deep_search_paths(query=str(action.get("query") or action.get("hint") or ""), roots=roots, max_results=int(action.get("max_results", 40) or 40), include_all_drives=_to_bool(action.get("include_all_drives", False), default=False), include_content=_to_bool(action.get("include_content", False), default=False), case_sensitive=_to_bool(action.get("case_sensitive", False), default=False), use_regex=_to_bool(action.get("use_regex", False), default=False), file_extensions=_to_str_list(action.get("file_extensions")) or None, max_scan_entries=int(action.get("max_scan_entries", 120000) or 120000), max_file_size_mb=int(action.get("max_file_size_mb", 5) or 5))
                return self._attach_mirror_context(result, mirror_context)
            if action_name == "analyze_path":
                mirror_context = self._get_mirror_context(self._mirror_query_from_action(action, action.get("path") or ""), limit=int(action.get("mirror_limit", 6) or 6)) if self._mirror_requested(action) else None
                result = self.analyze_path(path=str(action.get("path", "")), max_items=int(action.get("max_items", 3500) or 3500))
                return self._attach_mirror_context(result, mirror_context)
            if action_name == "create_path":
                return self.create_path(path=str(action.get("path", "")), kind=str(action.get("kind", "file")), content=str(action.get("content", "")), overwrite=_to_bool(action.get("overwrite", False), default=False))
            if action_name == "move_path":
                return self.move_path(src=str(action.get("src") or action.get("source") or ""), dst=str(action.get("dst") or action.get("destination") or ""), overwrite=_to_bool(action.get("overwrite", False), default=False))
            if action_name == "copy_path":
                return self.copy_path(src=str(action.get("src") or action.get("source") or ""), dst=str(action.get("dst") or action.get("destination") or ""), overwrite=_to_bool(action.get("overwrite", False), default=False))
            if action_name == "delete_path":
                return self.delete_path(path=str(action.get("path", "")), recursive=_to_bool(action.get("recursive", False), default=False), use_trash=_to_bool(action.get("use_trash", True), default=True))
            if action_name in {"search_file", "semantic_search_file"}:
                query = str(action.get("hint") or action.get("query") or action.get("mirror_query") or "")
                limit = int(action.get("max_results", 20) or 20)
                return self.search_mirror(query=query, limit=limit)
            if action_name == "read_file":
                mirror_context = self._get_mirror_context(self._mirror_query_from_action(action, action.get("path") or ""), limit=int(action.get("mirror_limit", 6) or 6)) if self._mirror_requested(action) else None
                result = self.read_file_text(file_path=str(action.get("path", "")), max_chars=int(action.get("max_chars", 12000) or 12000))
                return self._attach_mirror_context(result, mirror_context)
            if action_name in {"open_file", "open_path"}:
                mirror_context = self._get_mirror_context(self._mirror_query_from_action(action, action.get("path") or action.get("hint") or ""), limit=int(action.get("mirror_limit", 6) or 6)) if self._mirror_requested(action) else None
                result = self.open_file(file_path=str(action.get("path") or action.get("hint") or ""), resolve_by_hint=_to_bool(action.get("resolve_by_hint", True), default=True), roots=roots, include_content=_to_bool(action.get("include_content", True), default=True))
                return self._attach_mirror_context(result, mirror_context)
            
            # --- System/OS Interactions ---
            if action_name == "move_mouse":
                return self.move_mouse(x=int(action.get("x", 0)), y=int(action.get("y", 0)), duration=float(action.get("duration", 0.2)))
            if action_name == "click":
                return self.click(button=str(action.get("button", "left")), clicks=int(action.get("clicks", 1)))
            if action_name == "type_text":
                return self.type_text(text=str(action.get("text", "")), interval=float(action.get("interval", 0.01)))
            if action_name == "press_key":
                return self.press_key(str(action.get("key", "enter")))
            if action_name == "hotkey":
                keys = _to_str_list(action.get("keys"))
                if not keys: return _fail("Hotkey keys must be a non-empty list.", "invalid_hotkey_payload")
                return self.hotkey(keys)
            if action_name == "run_command":
                return self.run_system_command(command=str(action.get("command", "")), timeout=int(action.get("timeout", 25) or 25))
            if action_name == "open_interpreter":
                return self.run_open_interpreter(instruction=str(action.get("instruction") or action.get("prompt") or ""), timeout=int(action.get("timeout", 120) or 120))
            if action_name == "toggle_dark_mode":
                enable = action.get("enable")
                return self.toggle_dark_mode(None if enable is None else _to_bool(enable, default=False))
            
            # --- Communication Tools ---
            if action_name == "send_email":
                attachments = _to_str_list(action.get("attachments"))
                attachments.extend(_to_str_list(action.get("attachment")))
                file_hint = str(action.get("file_hint", "")).strip()
                if file_hint: attachments.append(file_hint)
                return self.send_email(to_email=str(action.get("to") or action.get("to_email") or ""), subject=str(action.get("subject", "")), body=str(action.get("body") or action.get("message") or ""), provider=str(action.get("provider", "gmail")), attachments=attachments, roots=roots, smtp_host=str(action.get("smtp_host", "")), smtp_port=int(action.get("smtp_port", 587) or 587), smtp_user=str(action.get("smtp_user", "")), smtp_password=str(action.get("smtp_password", "")), from_email=str(action.get("from") or action.get("from_email") or ""), send_now=_to_bool(action.get("send_now", False), default=False), use_outlook_desktop=_to_bool(action.get("use_outlook_desktop", True), default=True))
            if action_name == "draft_email_attachment":
                return self.draft_email_attachment(to_email=str(action.get("to") or action.get("to_email") or ""), subject=str(action.get("subject", "")), body=str(action.get("body", "")), file_hint=str(action.get("file_hint", "")), roots=roots, smtp_host=str(action.get("smtp_host", "")), smtp_port=int(action.get("smtp_port", 587) or 587), smtp_user=str(action.get("smtp_user", "")), smtp_password=str(action.get("smtp_password", "")), from_email=str(action.get("from") or action.get("from_email") or ""), send_now=_to_bool(action.get("send_now", False), default=False))
            if action_name == "send_telegram":
                return self.send_telegram(bot_token=str(action.get("bot_token") or action.get("token") or ""), chat_id=str(action.get("chat_id") or action.get("to") or ""), message=str(action.get("message") or action.get("text") or ""), file_path=str(action.get("file_path") or action.get("path") or ""), file_hint=str(action.get("file_hint") or action.get("attachment") or ""), roots=roots, disable_web_preview=_to_bool(action.get("disable_web_preview", False), default=False))
            if action_name == "send_whatsapp":
                return self.send_whatsapp(to_number=str(action.get("to") or action.get("to_number") or ""), message=str(action.get("message") or action.get("text") or ""), use_twilio=_to_bool(action.get("use_twilio", False), default=False), twilio_account_sid=str(action.get("twilio_account_sid") or action.get("account_sid") or ""), twilio_auth_token=str(action.get("twilio_auth_token") or action.get("auth_token") or ""), twilio_from=str(action.get("twilio_from") or action.get("from") or ""), media_url=str(action.get("media_url", "")))
            if action_name == "online_query":
                return self.connectivity.online_query(query=str(action.get("query", "")), max_results=int(action.get("max_results", 5) or 5))
            
            # --- Office Document Creation ---
            if action_name in {"create_word_doc", "create_pdf_report", "create_excel_report", "analyze_excel", "analyze_word_doc", "analyze_pdf"}:
                return execute_office_tool(action_name, action)
            
            return _fail("Unknown tool action.", f"unsupported_action: {action_name}")
        except Exception as exc:
            return _fail("Tool dispatcher failed.", str(exc))

# --- Single Global Task Bridge Instance ---
_TASK_BRIDGE = UnifiedTaskBridge()

def run_tool_action(action: Dict[str, Any]) -> Dict[str, Any]:
    return _TASK_BRIDGE.execute_action(action)

def search_files_by_hint(hint: str, roots: Optional[List[str]] = None, max_results: int = 20) -> Dict[str, Any]:
    return _TASK_BRIDGE.search_files_by_hint(hint=hint, roots=roots, max_results=max_results)

def read_file_text(file_path: str, max_chars: int = 12000) -> Dict[str, Any]:
    return _TASK_BRIDGE.read_file_text(file_path=file_path, max_chars=max_chars)

def open_file(file_path: str) -> Dict[str, Any]:
    return _TASK_BRIDGE.open_file(file_path=file_path)

def move_mouse(x: int, y: int, duration: float = 0.2) -> Dict[str, Any]:
    return _TASK_BRIDGE.move_mouse(x=x, y=y, duration=duration)

def click(button: str = "left", clicks: int = 1) -> Dict[str, Any]:
    return _TASK_BRIDGE.click(button=button, clicks=clicks)

def type_text(text: str, interval: float = 0.01) -> Dict[str, Any]:
    return _TASK_BRIDGE.type_text(text=text, interval=interval)

def press_key(key: str) -> Dict[str, Any]:
    return _TASK_BRIDGE.press_key(key=key)

def hotkey(keys: List[str]) -> Dict[str, Any]:
    return _TASK_BRIDGE.hotkey(keys=keys)


# =============================================================================
# 6. NLP ACTION COMMAND HANDLERS
# =============================================================================

class ExcelCommandHandler:
    def __init__(self):
        self.default_extension = ".xlsx"

    def _load_or_create_workbook(self, path):
        if os.path.exists(path): return load_workbook(path)
        return Workbook()

    def _pick_sheet(self, workbook, sheet_name):
        if not sheet_name: return workbook.active
        target_sheet = sheet_name.strip().strip('"').strip("'")
        if target_sheet in workbook.sheetnames: return workbook[target_sheet]
        return workbook.create_sheet(target_sheet)

    def _to_value(self, raw_value):
        value = raw_value.strip()
        if (value.startswith('"') and value.endswith('"')) or (value.startswith("'") and value.endswith("'")): return value[1:-1]
        if value.lower() == "true": return True
        if value.lower() == "false": return False
        if re.fullmatch(r"-?\d+", value): return int(value)
        try: return float(value)
        except ValueError: pass
        return value

    def _ensure_formula_prefix(self, formula_text):
        cleaned_formula = formula_text.strip()
        return cleaned_formula if cleaned_formula.startswith("=") else "=" + cleaned_formula

    def _create_random_table_with_chart(self, file_name, rows=10, cols=10):
        file_path = _normalize_office_path(file_name, self.default_extension)
        workbook = self._load_or_create_workbook(file_path)
        sheet = self._pick_sheet(workbook, "RandomData")

        if sheet.max_row > 0: sheet.delete_rows(1, sheet.max_row)
        if sheet.max_column > 0: sheet.delete_cols(1, sheet.max_column)
        sheet._charts = []

        for row in range(1, rows + 1):
            for col in range(1, cols + 1):
                sheet.cell(row=row, column=col, value=random.randint(10, 99))

        chart = LineChart()
        chart.title = f"Random Data {rows}x{cols}"
        chart.style = 10
        chart.y_axis.title = "Value"
        chart.x_axis.title = "Row"

        data_ref = Reference(sheet, min_col=1, min_row=1, max_col=cols, max_row=rows)
        category_ref = Reference(sheet, min_col=1, min_row=1, max_row=rows)
        chart.add_data(data_ref, titles_from_data=False)
        chart.set_categories(category_ref)

        chart_anchor = f"{get_column_letter(cols + 2)}2"
        sheet.add_chart(chart, chart_anchor)
        workbook.save(file_path)
        print(f"[ACTION][EXCEL] Random {rows}x{cols} table + chart created in {file_path}")

        try:
            if _open_with_default_app(file_path):
                print("[ACTION][EXCEL] Opened workbook in available spreadsheet software.")
            else:
                print("[ACTION][EXCEL] Workbook created, but opening the file failed.")
        except Exception as e:
            print(f"[ACTION][EXCEL] Workbook created but auto-open failed: {e}")

    def handle(self, text):
        random_nl_match = re.fullmatch(r"(?:create|make)\s+an?\s+excel\s+workbook\s+with\s+(\d+)x(\d+)\s+table\s+with\s+random\s+data\s+and\s+then\s+create\s+(?:the\s+)?graph\s+from\s+it(?:\s+in\s+available\s+software)?(?:\s+in\s+(.+?))?\s*$", text, flags=re.IGNORECASE)
        if random_nl_match:
            rows, cols, file_name = random_nl_match.groups()
            rows, cols = int(rows), int(cols)
            if rows > 100 or cols > 50 or rows <= 0 or cols <= 0:
                print("[ACTION][EXCEL] Table size out of supported range. Use 1..100 rows and 1..50 columns.")
                return True
            self._create_random_table_with_chart(file_name or "random_chart_demo.xlsx", rows=rows, cols=cols)
            return True

        random_cmd_match = re.fullmatch(r"excel\s+random\s+table\s+(\d+)x(\d+)\s+with\s+graph(?:\s+in\s+(.+?))?\s*$", text, flags=re.IGNORECASE)
        if random_cmd_match:
            rows, cols, file_name = random_cmd_match.groups()
            rows, cols = int(rows), int(cols)
            if rows > 100 or cols > 50 or rows <= 0 or cols <= 0:
                print("[ACTION][EXCEL] Table size out of supported range. Use 1..100 rows and 1..50 columns.")
                return True
            self._create_random_table_with_chart(file_name or "random_chart_demo.xlsx", rows=rows, cols=cols)
            return True

        demo_file_match = re.fullmatch(r"excel\s+demo\s+table\s+graph(?:\s+in\s+(.+?))?\s*$", text, flags=re.IGNORECASE)
        if demo_file_match:
            self._create_random_table_with_chart(demo_file_match.group(1) or "random_chart_demo.xlsx", rows=10, cols=10)
            return True

        open_match = re.fullmatch(r"excel\s+(?:create|open)\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if open_match:
            file_path = _normalize_office_path(open_match.group(1), self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] Ready: {file_path}")
            return True

        create_sheet_match = re.fullmatch(r"excel\s+create\s+sheet\s+(.+?)\s+in\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if create_sheet_match:
            sheet_name, file_name = create_sheet_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            target_sheet = sheet_name.strip().strip('"').strip("'")
            if target_sheet not in workbook.sheetnames:
                workbook.create_sheet(target_sheet)
                workbook.save(file_path)
                print(f"[ACTION][EXCEL] Sheet '{target_sheet}' created in {file_path}")
            else:
                print(f"[ACTION][EXCEL] Sheet '{target_sheet}' already exists in {file_path}")
            return True

        list_sheets_match = re.fullmatch(r"excel\s+list\s+sheets\s+in\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if list_sheets_match:
            file_path = _normalize_office_path(list_sheets_match.group(1), self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            print(f"[ACTION][EXCEL] Sheets in {file_path}: {', '.join(workbook.sheetnames)}")
            return True

        set_match = re.fullmatch(r"excel\s+set\s+([a-zA-Z]+\d+)\s+to\s+(.+?)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if set_match:
            cell, value_text, file_name, sheet_name = set_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            sheet[cell.upper()] = self._to_value(value_text)
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] {cell.upper()} updated in {file_path}")
            return True

        get_match = re.fullmatch(r"excel\s+get\s+([a-zA-Z]+\d+)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if get_match:
            cell, file_name, sheet_name = get_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            _ = sheet[cell.upper()].value
            print(f"[ACTION][EXCEL] {cell.upper()} retrieved from {file_path}. Content hidden by policy.")
            return True

        row_match = re.fullmatch(r"excel\s+add\s+row\s+(.+?)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if row_match:
            row_values_text, file_name, sheet_name = row_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            try:
                parsed_row = next(csv.reader(StringIO(row_values_text)), [])
            except csv.Error:
                print("[ACTION][EXCEL] Row input is malformed. Use comma-separated values, e.g. apple,10,done.")
                return True
            if not parsed_row or all(not item.strip() for item in parsed_row):
                print("[ACTION][EXCEL] No valid row values provided. Use comma-separated values.")
                return True
            sheet.append([self._to_value(item.strip()) for item in parsed_row])
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] Row added in {file_path}")
            return True

        delete_row_match = re.fullmatch(r"excel\s+delete\s+row\s+(\d+)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if delete_row_match:
            row_number, file_name, sheet_name = delete_row_match.groups()
            row_number = int(row_number)
            if row_number <= 0:
                print("[ACTION][EXCEL] Row number must be 1 or greater.")
                return True
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            sheet.delete_rows(row_number, 1)
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] Row {row_number} deleted in {file_path}")
            return True

        delete_col_match = re.fullmatch(r"excel\s+delete\s+column\s+([a-zA-Z]+)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if delete_col_match:
            col_letters, file_name, sheet_name = delete_col_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            col_letters = col_letters.upper()
            col_index = 0
            for ch in col_letters: col_index = col_index * 26 + (ord(ch) - ord("A") + 1)
            sheet.delete_cols(col_index, 1)
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] Column {col_letters} deleted in {file_path}")
            return True

        sum_match = re.fullmatch(r"excel\s+sum\s+([a-zA-Z]+\d+:[a-zA-Z]+\d+)\s+in\s+(.+?)\s+to\s+([a-zA-Z]+\d+)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if sum_match:
            source_range, file_name, target_cell, sheet_name = sum_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            sheet[target_cell.upper()] = self._ensure_formula_prefix(f"SUM({source_range.upper()})")
            workbook.save(file_path)
            print(f"[ACTION][SUM] Sum {source_range.upper()} -> {target_cell.upper()} in {file_path}")
            return True

        formula_match = re.fullmatch(r"excel\s+formula\s+([a-zA-Z]+\d+)\s*=\s*(.+?)\s+in\s+(.+?)(?:\s+sheet\s+(.+))?", text, flags=re.IGNORECASE)
        if formula_match:
            target_cell, formula_text, file_name, sheet_name = formula_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            workbook = self._load_or_create_workbook(file_path)
            sheet = self._pick_sheet(workbook, sheet_name)
            sheet[target_cell.upper()] = self._ensure_formula_prefix(formula_text)
            workbook.save(file_path)
            print(f"[ACTION][EXCEL] Formula set in {target_cell.upper()} for {file_path}")
            return True

        if re.fullmatch(r"excel\s+help", text, flags=re.IGNORECASE):
            print("[ACTION][EXCEL] Commands:")
            for command in EXCEL_HELP_COMMANDS: print(f"- {command}")
            return True

        return False


class WordCommandHandler:
    def __init__(self):
        self.default_extension = ".docx"

    def _load_or_create_doc(self, path):
        if not DOCX_AVAILABLE: return None
        if os.path.exists(path): return Document(path)
        return Document()

    def handle(self, text):
        if not DOCX_AVAILABLE:
            if text.lower().startswith("word "): print("[ACTION][WORD] python-docx is not installed. Run: pip install python-docx")
            return text.lower().startswith("word ")

        create_match = re.fullmatch(r"word\s+(?:create|open)\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if create_match:
            file_path = _normalize_office_path(create_match.group(1), self.default_extension)
            doc = self._load_or_create_doc(file_path)
            doc.save(file_path)
            print(f"[ACTION][WORD] Ready: {file_path}")
            return True

        paragraph_match = re.fullmatch(r"word\s+add\s+paragraph\s+(.+?)\s+in\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if paragraph_match:
            paragraph_text, file_name = paragraph_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            doc = self._load_or_create_doc(file_path)
            content = paragraph_text.strip().strip('"').strip("'")
            doc.add_paragraph(content)
            doc.save(file_path)
            print(f"[ACTION][WORD] Paragraph added in {file_path}")
            return True

        heading_match = re.fullmatch(r"word\s+add\s+heading\s+(.+?)(?:\s+level\s+([1-6]))?\s+in\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if heading_match:
            heading_text, level_text, file_name = heading_match.groups()
            level = int(level_text) if level_text else 1
            file_path = _normalize_office_path(file_name, self.default_extension)
            doc = self._load_or_create_doc(file_path)
            content = heading_text.strip().strip('"').strip("'")
            doc.add_heading(content, level=level)
            doc.save(file_path)
            print(f"[ACTION][WORD] Heading (level {level}) added in {file_path}")
            return True

        read_match = re.fullmatch(r"word\s+read\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if read_match:
            file_path = _normalize_office_path(read_match.group(1), self.default_extension)
            if not os.path.exists(file_path):
                print(f"[ACTION][WORD] File not found: {file_path}")
                return True
            doc = self._load_or_create_doc(file_path)
            print(f"[ACTION][WORD] Document read from {file_path}. Content hidden by policy.")
            return True

        if re.fullmatch(r"word\s+help", text, flags=re.IGNORECASE):
            print("[ACTION][WORD] Commands:")
            for command in WORD_HELP_COMMANDS: print(f"- {command}")
            return True

        return False


class PowerPointCommandHandler:
    def __init__(self):
        self.default_extension = ".pptx"

    def _load_or_create_presentation(self, path):
        if not POWERPOINT_AVAILABLE: return None
        if os.path.exists(path): return Presentation(path)
        return Presentation()

    def handle(self, text):
        if not POWERPOINT_AVAILABLE:
            if text.lower().startswith("powerpoint "): print("[ACTION][PPT] python-pptx is not installed. Run: pip install python-pptx")
            return text.lower().startswith("powerpoint ")

        create_match = re.fullmatch(r"powerpoint\s+(?:create|open)\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if create_match:
            file_path = _normalize_office_path(create_match.group(1), self.default_extension)
            prs = self._load_or_create_presentation(file_path)
            prs.save(file_path)
            print(f"[ACTION][PPT] Ready: {file_path}")
            return True

        slide_match = re.fullmatch(r"powerpoint\s+add\s+slide\s+title\s+(.+?)\s+content\s+(.+?)\s+in\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if slide_match:
            title_text, content_text, file_name = slide_match.groups()
            file_path = _normalize_office_path(file_name, self.default_extension)
            prs = self._load_or_create_presentation(file_path)
            layout = prs.slide_layouts[1] if len(prs.slide_layouts) > 1 else prs.slide_layouts[0]
            slide = prs.slides.add_slide(layout)

            title_box = getattr(slide.shapes, "title", None)
            if title_box: title_box.text = title_text.strip().strip('"').strip("'")
            if len(slide.placeholders) > 1: slide.placeholders[1].text = content_text.strip().strip('"').strip("'")

            prs.save(file_path)
            print(f"[ACTION][PPT] Slide added in {file_path}")
            return True

        launch_match = re.fullmatch(r"powerpoint\s+launch\s+(.+?)\s*$", text, flags=re.IGNORECASE)
        if launch_match:
            file_path = _normalize_office_path(launch_match.group(1), self.default_extension)
            if not os.path.exists(file_path):
                print(f"[ACTION][PPT] File not found: {file_path}")
                return True
            try:
                if _open_with_default_app(file_path): print(f"[ACTION][PPT] Opened {file_path}")
                else: print(f"[ACTION][PPT] Failed to open {file_path}")
            except Exception as e:
                print(f"[ACTION][PPT] Failed to open: {e}")
            return True

        if re.fullmatch(r"powerpoint\s+help", text, flags=re.IGNORECASE):
            print("[ACTION][PPT] Commands:")
            for command in POWERPOINT_HELP_COMMANDS: print(f"- {command}")
            return True

        return False


class ActionHandler:
    # --- PRE-COMPILED REGEX OPTIMIZATION ---
    NATURAL_SEARCH_RX = re.compile(r"\b(?:search|find|locate|look\s+for)\b.*?\b(?:file|files|folder|folders|document|documents|path)\b\s+(.+)$", re.IGNORECASE)
    NATURAL_OPEN_RX = re.compile(r"\bopen\b.*?\b(?:file|folder|document|path)\b\s+(.+)$", re.IGNORECASE)
    NATURAL_LAUNCH_RX = re.compile(r"\b(?:open|launch|start|run)\b\s+(.+)$", re.IGNORECASE)
    NATURAL_CLOSE_RX = re.compile(r"\b(?:close|quit|exit|kill|terminate)\b\s+(.+)$", re.IGNORECASE)
    
    CMD_FILES_ROOTS_RX = re.compile(r"files\s+roots\s*", re.IGNORECASE)
    CMD_FILES_LIST_RX = re.compile(r"files\s+list\s+(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_SEARCH_RX = re.compile(r"files\s+deep\s+search\s+(.+?)(?:\s+in\s+(.+))?\s*$", re.IGNORECASE)
    CMD_FILES_ANALYZE_RX = re.compile(r"files\s+analyze\s+(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_CREATE_RX = re.compile(r"files\s+create\s+file\s+(.+?)(?:\s+content\s+(.+))?\s*$", re.IGNORECASE)
    CMD_FILES_MKDIR_RX = re.compile(r"files\s+create\s+(?:folder|dir|directory)\s+(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_MOVE_RX = re.compile(r"files\s+move\s+(.+?)\s*->\s*(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_COPY_RX = re.compile(r"files\s+copy\s+(.+?)\s*->\s*(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_DEL_RX = re.compile(r"files\s+(?:delete|remove)\s+(.+?)\s*$", re.IGNORECASE)
    CMD_FILES_OPEN_RX = re.compile(r"files\s+open\s+(.+?)\s*$", re.IGNORECASE)
    
    CMD_SOFT_OPEN_RX = re.compile(r"software\s+open\s+(.+?)\s*$", re.IGNORECASE)
    CMD_SOFT_CLOSE_RX = re.compile(r"software\s+close\s+(.+?)\s*$", re.IGNORECASE)
    CMD_SOFT_RUN_RX = re.compile(r"software\s+running\s*", re.IGNORECASE)
    CMD_SVC_OPEN_RX = re.compile(r"service\s+open\s+(.+?)\s*$", re.IGNORECASE)
    
    CMD_EMAIL_RX = re.compile(r"email\s+(draft|send)\s+to\s+(.+?)\s+subject\s+(.+?)\s+body\s+(.+?)(?:\s+attach\s+(.+?))?(?:\s+provider\s+(gmail|outlook|custom))?\s*$", re.IGNORECASE)
    CMD_TG_SEND_RX = re.compile(r"telegram\s+send\s+to\s+(.+?)\s+token\s+(.+?)\s+message\s+(.+)\s*$", re.IGNORECASE)
    CMD_TG_FILE_RX = re.compile(r"telegram\s+file\s+(.+?)\s+to\s+(.+?)\s+token\s+(.+?)(?:\s+caption\s+(.+))?\s*$", re.IGNORECASE)
    CMD_WA_SEND_RX = re.compile(r"whatsapp\s+send\s+to\s+(.+?)\s+message\s+(.+)\s*$", re.IGNORECASE)

    CMD_WEB_RESEARCH_RX = re.compile(r"(?:web\s+research|research\s+web|research)\s+(.+)", re.IGNORECASE)
    CMD_OPEN_WEB_RX = re.compile(r"(?:open\s+website|browse)\s+(.+)", re.IGNORECASE)
    CMD_SEARCH_WEB_RX = re.compile(r"search\s+web\s+(.+)", re.IGNORECASE)
    CMD_SAVE_CLIP_RX = re.compile(r"save\s+clipboard\s+to\s+rad\s+as\s+(.+)", re.IGNORECASE)
    CMD_COPY_TEXT_RX = re.compile(r"(?:copy\s+selected\s+text|copy\s+now)", re.IGNORECASE)
    CMD_PASTE_TEXT_RX = re.compile(r"(?:paste\s+clipboard|paste\s+now)", re.IGNORECASE)
    CMD_SYS_CHECK_RX = re.compile(r"(?:system check|check system)", re.IGNORECASE)
    CMD_MALWARE_SCAN_RX = re.compile(r"(?:malware scan|scan for malware|run malware scan|security quick scan)", re.IGNORECASE)

    @staticmethod
    def get_supported_command_sections():
        return {
            "General task commands": list(GENERAL_TASK_COMMANDS),
            "File system commands": list(FILE_SYSTEM_COMMANDS),
            "Software control commands": list(SOFTWARE_COMMANDS),
            "Communication commands": list(COMMUNICATION_COMMANDS),
            "Clipboard commands": list(CLIPBOARD_COMMANDS),
            "System commands": list(SYSTEM_COMMANDS),
            "Office quick help": list(OFFICE_HELP_COMMANDS),
            "Excel commands": list(EXCEL_HELP_COMMANDS),
            "Word commands": list(WORD_HELP_COMMANDS),
            "PowerPoint commands": list(POWERPOINT_HELP_COMMANDS),
            "Assistant JSON command format": list(ASSISTANT_JSON_ACTIONS),
            "Assistant tool actions (action field)": list(ASSISTANT_TOOL_ACTIONS),
            "Assistant tool call examples": list(ASSISTANT_TOOL_EXAMPLES),
        }

    def __init__(self, db=None, context_provider=None):
        self._explicit_db = db
        self._db_instance = None
        self.context_provider = context_provider
        self.excel = ExcelCommandHandler()
        self.word = WordCommandHandler()
        self.powerpoint = PowerPointCommandHandler()
        action_cfg = CONFIG.get("actions", {})
        self.safe_mode = bool(action_cfg.get("safe_mode", False))
        self.allow_legacy_text_commands = bool(action_cfg.get("allow_legacy_text_commands", True))

    @property
    def db(self):
        if self._explicit_db is not None: return self._explicit_db
        if self._db_instance is None:
            try:
                from aiassistant.infra.db.database_manager import DatabaseManager
                db_file_path = str(CONFIG.get("paths", {}).get("db_path", "cache/assistant_sessions.db"))
                self._db_instance = DatabaseManager(db_path=db_file_path)
            except Exception as e:
                print(f"[ACTION][INIT] Lazy database fallback assignment failed: {e}")
                self._db_instance = None
        return self._db_instance

    def execute_and_collect(self, text):
        if not text: return ""
        buffer = io.StringIO()
        with contextlib.redirect_stdout(buffer):
            try:
                self.execute(text)
            except Exception as e:
                print(f"[ACTION ERROR] {e}")
        return buffer.getvalue().strip()

    def _looks_like_action_command(self, text):
        if not text: return False
        lowered = text.strip().lower()
        prefixes = (
            "files ", "software ", "service ", "email ", "telegram ", "whatsapp ",
            "search file ", "find file ", "locate file ", "open file ", "open folder ",
            "open document ", "open ", "close ", "play ", "write ", "note ", "type ",
            "take a note ", "search web ", "open website ", "browse ", "excel ",
            "word ", "powerpoint ", "research ", "web research ", "system check",
            "check system", "scan apps", "update apps", "copy selected text",
            "copy now", "paste clipboard", "paste now", "save clipboard to rad as ",
            "malware scan", "scan for malware", "run malware scan", "security quick scan",
            "schedule ", "meeting "
        )
        if lowered.startswith(prefixes): return True
        if re.search(r"\b(?:open|close)\b\s+[a-z0-9]", lowered): return True
        if re.search(r"\b(?:open|search|find|locate|look\s+for)\b.*\b(?:file|files|folder|folders|document|documents|path)\b", lowered): return True
        return any(token in lowered for token in ("volume up", "volume down", "mute", "unmute"))

    @staticmethod
    def _looks_like_file_hint(raw_text):
        clean = str(raw_text or "").strip().lower()
        if not clean: return False
        if re.search(r"^[a-z]:\\", clean): return True
        if "\\" in clean or "/" in clean: return True
        if re.search(r"\.[a-z0-9]{1,6}\b", clean): return True
        file_tokens = (" file", "files ", "folder", "document", "directory", "path", "report", "notes", "draft", "summary", "invoice", "project")
        return any(token in clean for token in file_tokens)

    def _print_system_check(self):
        checks = {
            "Python": True,
            "Reasoning Server Module": importlib.util.find_spec("aiassistant.backend.server_reasoning") is not None,
            "Voice Server Module": importlib.util.find_spec("aiassistant.backend.server_voice") is not None,
            "Piper Folder": os.path.isdir("piper"),
            "Models Folder": os.path.isdir("models"),
            "RVC Models Folder": os.path.isdir("rvc"),
            "Internet Library (requests)": REQUESTS_AVAILABLE,
            "Clipboard Library (pyperclip)": CLIPBOARD_AVAILABLE,
        }
        print("[ACTION][SYSTEM] Quick Check:")
        for name, is_ok in checks.items(): print(f"- {name}: {'OK' if is_ok else 'MISSING'}")

    def _run_quick_malware_scan(self):
        print("[ACTION][SECURITY] Starting Windows Defender quick scan...")
        try:
            subprocess.run(["powershell", "-NoProfile", "-Command", "Start-MpScan -ScanType QuickScan"], check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            print("[ACTION][SECURITY] Quick scan request sent to Windows Defender.")
        except Exception as e: print(f"[ACTION][SECURITY] Could not start Defender scan: {e}")

    def _extract_key_points(self, text, max_points=5):
        if not text: return []
        chunks = re.split(r"(?<=[.!?])\s+", text)
        points = []
        for chunk in chunks:
            cleaned = re.sub(r"\s+", " ", chunk).strip(" -\n\t")
            if len(cleaned) < 30 or cleaned in points: continue
            points.append(cleaned)
            if len(points) >= max_points: break
        return points

    def _store_web_points_to_rad(self, query, points, source_url=""):
        if not self.db or not hasattr(self.db, "add_rad_data_if_new"): return 0
        slug = re.sub(r"[^a-z0-9]+", "_", query.lower()).strip("_")[:32] or "query"
        stored_count = 0
        for idx, point in enumerate(points, start=1):
            if self.db.add_rad_data_if_new("web_fact", f"web_{slug}_{idx}", point, 0.78): stored_count += 1
        if source_url and self.db.add_rad_data_if_new("web_source", f"source_{slug}", source_url, 0.95): stored_count += 1
        return stored_count

    def _fetch_web_text(self, url):
        if not REQUESTS_AVAILABLE: return ""
        response = requests.get(url, headers={"User-Agent": "Mozilla/5.0 (compatible; MARIE/1.0)"}, timeout=12)
        response.raise_for_status()
        html_text = response.text
        if BS4_AVAILABLE:
            soup = BeautifulSoup(html_text, "html.parser")
            paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
            return re.sub(r"\s+", " ", " ".join(paragraphs)).strip()

        no_scripts = re.sub(r"<script[\s\S]*?</script>", " ", html_text, flags=re.IGNORECASE)
        no_styles = re.sub(r"<style[\s\S]*?</style>", " ", no_scripts, flags=re.IGNORECASE)
        plain = re.sub(r"<[^>]+>", " ", no_styles)
        return re.sub(r"\s+", " ", plain).strip()

    def _research_and_store_web_data(self, query):
        if not REQUESTS_AVAILABLE:
            print("[ACTION][WEB] requests is not installed. Run: pip install requests")
            return

        query = query.strip()
        if not query:
            print("[ACTION][WEB] Missing query. Example: web research latest AI trends")
            return

        source_url = ""
        combined_text = ""

        try:
            if query.startswith(("http://", "https://")):
                source_url = query
                combined_text = self._fetch_web_text(source_url)
            else:
                ddg_url = "https://api.duckduckgo.com/"
                response = requests.get(ddg_url, params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1}, timeout=10)
                response.raise_for_status()
                data = response.json()

                abstract = data.get("AbstractText", "")
                source_url = data.get("AbstractURL", "")
                related_lines = []
                for item in data.get("RelatedTopics", [])[:10]:
                    if isinstance(item, dict) and item.get("Text"): related_lines.append(item["Text"])
                    elif isinstance(item, dict) and item.get("Topics"):
                        for sub in item.get("Topics", [])[:5]:
                            if isinstance(sub, dict) and sub.get("Text"): related_lines.append(sub["Text"])

                combined_text = " ".join([abstract] + related_lines)

                if source_url:
                    fetched_page_text = self._fetch_web_text(source_url)
                    if fetched_page_text: combined_text = f"{combined_text} {fetched_page_text}"

            points = self._extract_key_points(combined_text, max_points=6)
            if not points:
                print("[ACTION][WEB] Could not extract useful points from the web result.")
                return

            print(f"[ACTION][WEB] Key points for '{query}':")
            for idx, point in enumerate(points, start=1): print(f"{idx}. {point}")

            stored_count = self._store_web_points_to_rad(query, points, source_url)
            if stored_count > 0: print(f"[ACTION][RAD] Stored {stored_count} web-derived memory items.")
            if CLIPBOARD_AVAILABLE:
                pyperclip.copy("\n".join(points))
                print("[ACTION][WEB] Copied key points to clipboard.")
        except Exception as e:
            print(f"[ACTION][WEB] Research failed: {e}")

    def _open_text_editor(self):
        try:
            if IS_WINDOWS:
                os.system("start notepad")
                return
            if SYSTEM_NAME == "darwin":
                subprocess.Popen(["open", "-a", "TextEdit"])
                return
            for editor in ("gedit", "kate", "mousepad", "xed"):
                if shutil.which(editor):
                    subprocess.Popen([editor])
                    return
        except Exception as e: print(f"[ACTION] Could not open text editor: {e}")

    def _extract_json_action(self, text):
        if not text: return None
        candidates = []
        stripped = text.strip()
        if stripped.startswith("{") and stripped.endswith("}"): candidates.append(stripped)

        fenced = re.findall(r"```json\s*(\{[\s\S]*?\})\s*```", text, flags=re.IGNORECASE)
        candidates.extend(fenced)

        loose = re.findall(r"(\{\s*\"action\"[\s\S]*?\})", text)
        candidates.extend(loose)

        for candidate in candidates:
            try: obj = json.loads(candidate)
            except Exception: continue
            if not isinstance(obj, dict): continue

            action = str(obj.get("action", "")).strip().lower()
            target = str(obj.get("target", "")).strip()
            value = obj.get("value", "")
            if action in {"open", "close", "search_web", "open_website", "volume", "write_note", "play"}:
                return {"action": action, "target": target, "value": value}

        return None

    def _print_tool_bridge_result(self, result):
        if not isinstance(result, dict):
            print(f"[ACTION][TOOLS] {result}")
            return False

        ok = bool(result.get("success", False))
        prefix = "[ACTION][TOOLS]" if ok else "[ACTION][TOOLS][ERROR]"
        message = str(result.get("message", "")).strip()
        if message and not "found 0 lexical file match" in message.lower():
            print(f"{prefix} {message}")

        payload = result.get("data")
        if TOOL_BRIDGE_VERBOSE and payload not in (None, [], {}, ""):
            try: text = json.dumps(payload, ensure_ascii=True)
            except Exception: text = str(payload)
            if len(text) > 1800: text = text[:1800] + "..."
            print(f"{prefix} Data: {text}")

        error = str(result.get("error", "")).strip()
        if error and not ok: print(f"{prefix} {error}")
        return ok

    def _run_tool_bridge_action(self, action):
        try: return self._print_tool_bridge_result(run_tool_action(action))
        except Exception as e:
            print(f"[ACTION][TOOLS][ERROR] {e}")
            return False

    def _handle_extended_tool_commands(self, raw_text):
        if not raw_text: return False
        lowered = raw_text.strip().lower()

        if match := self.NATURAL_SEARCH_RX.search(raw_text):
            if "website" not in lowered and "search web" not in lowered:
                hint = match.group(1).strip().strip('"').strip("'")
                if hint:
                    self._run_tool_bridge_action({"action": "search_file", "hint": hint, "max_results": 20, "include_content": True})
                    return True

        if match := self.NATURAL_OPEN_RX.search(raw_text):
            if "website" not in lowered:
                target_hint = match.group(1).strip().strip('"').strip("'")
                if target_hint:
                    self._run_tool_bridge_action({"action": "open_file", "path": target_hint, "resolve_by_hint": True, "include_content": True})
                    return True

        if match := self.NATURAL_LAUNCH_RX.search(raw_text):
            if not re.search(r"\b(file|files|folder|folders|document|documents|path)\b", lowered):
                app_target = match.group(1).strip().strip('"').strip("'")
                app_target = re.sub(r"\b(?:please|now|for\s+me|thanks|thank\s+you)\b", " ", app_target, flags=re.IGNORECASE)
                app_target = re.sub(r"\b(?:app|application|software|program)\b", " ", app_target, flags=re.IGNORECASE)
                app_target = re.sub(r"^(?:the|a|an)\s+", "", app_target, flags=re.IGNORECASE)
                app_target = re.sub(r"\s+", " ", app_target).strip()
                if app_target and "website" not in app_target.lower():
                    if re.search(r"(?:https?://|www\.|\.[a-z]{2,6}(?:/|$))", app_target, flags=re.IGNORECASE):
                        url = app_target if app_target.startswith("http") else f"https://{app_target}"
                        webbrowser.open(url)
                    else:
                        self._run_tool_bridge_action({"action": "launch_application", "app": app_target})
                    return True

        if match := self.NATURAL_CLOSE_RX.search(raw_text):
            if not re.search(r"\b(file|files|folder|folders|document|documents|path)\b", lowered):
                close_target = match.group(1).strip().strip('"').strip("'")
                close_target = re.sub(r"\b(?:please|now|for\s+me|thanks|thank\s+you)\b", " ", close_target, flags=re.IGNORECASE)
                close_target = re.sub(r"\b(?:app|application|software|program)\b", " ", close_target, flags=re.IGNORECASE)
                close_target = re.sub(r"^(?:the|a|an)\s+", "", close_target, flags=re.IGNORECASE)
                close_target = re.sub(r"\s+", " ", close_target).strip()
                if close_target:
                    self._run_tool_bridge_action({"action": "close_application", "app": close_target})
                    return True

        if self.CMD_FILES_ROOTS_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "list_system_roots"})
        if match := self.CMD_FILES_LIST_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "list_directory", "path": match.group(1).strip(), "max_results": 120})
        
        if match := self.CMD_FILES_SEARCH_RX.fullmatch(raw_text):
            query, root = match.groups()
            action = {"action": "deep_search", "query": query.strip(), "max_results": 40, "include_content": True}
            if root: action["roots"] = [root.strip()]
            else: action["include_all_drives"] = True
            return self._run_tool_bridge_action(action)

        if match := self.CMD_FILES_ANALYZE_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "analyze_path", "path": match.group(1).strip()})
        if match := self.CMD_FILES_CREATE_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "create_path", "path": match.group(1).strip(), "kind": "file", "content": (match.group(2) or "").strip(), "overwrite": False})
        if match := self.CMD_FILES_MKDIR_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "create_path", "path": match.group(1).strip(), "kind": "directory"})
        if match := self.CMD_FILES_MOVE_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "move_path", "src": match.group(1).strip(), "dst": match.group(2).strip()})
        if match := self.CMD_FILES_COPY_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "copy_path", "src": match.group(1).strip(), "dst": match.group(2).strip()})
        if match := self.CMD_FILES_DEL_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "delete_path", "path": match.group(1).strip(), "recursive": True, "use_trash": True})
        if match := self.CMD_FILES_OPEN_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "open_file", "path": match.group(1).strip(), "resolve_by_hint": True, "include_content": True})
        
        if match := self.CMD_SOFT_OPEN_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "launch_application", "app": match.group(1).strip()})
        if match := self.CMD_SOFT_CLOSE_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "close_application", "app": match.group(1).strip()})
        if self.CMD_SOFT_RUN_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "list_running_apps", "max_results": 120})
        if match := self.CMD_SVC_OPEN_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "open_service", "service": match.group(1).strip()})

        if match := self.CMD_EMAIL_RX.fullmatch(raw_text):
            mode, to_email, subject, body, attachment, provider = match.groups()
            payload = {"action": "send_email", "to": to_email.strip(), "subject": subject.strip(), "body": body.strip(), "provider": (provider or ("outlook" if IS_WINDOWS else "gmail")).strip(), "send_now": mode.strip().lower() == "send"}
            if attachment: payload["attachments"] = [attachment.strip()]
            return self._run_tool_bridge_action(payload)

        if match := self.CMD_TG_SEND_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "send_telegram", "chat_id": match.group(1).strip(), "bot_token": match.group(2).strip(), "message": match.group(3).strip()})
        if match := self.CMD_TG_FILE_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "send_telegram", "chat_id": match.group(2).strip(), "bot_token": match.group(3).strip(), "file_hint": match.group(1).strip(), "message": (match.group(4) or "").strip()})
        if match := self.CMD_WA_SEND_RX.fullmatch(raw_text): return self._run_tool_bridge_action({"action": "send_whatsapp", "to": match.group(1).strip(), "message": match.group(2).strip()})

        return False
    
    def _run_meeting_automation(self, subject: str, start_phrase: str, duration: int, attendees: list):
        """Helper method to clean dates and pass parameters to background Outlook automation."""
        import threading
        import time
        from datetime import datetime, timedelta
        
        def run_outlook():
            import pythoncom
            pythoncom.CoInitialize()
            try:
                import win32com.client as win32
                outlook = win32.Dispatch("Outlook.Application")
                appt = outlook.CreateItem(1) # olAppointmentItem
                
                # Simple parser for natural expressions
                now = datetime.now()
                target_date = now + timedelta(days=1) if "tomorrow" in start_phrase.lower() else now
                
                hour, minute = 10, 0
                digits = re.findall(r"\d+", start_phrase)
                if digits:
                    hour = int(digits[0])
                    if len(digits) > 1: minute = int(digits[1])
                    if "pm" in start_phrase.lower() and hour < 12: hour += 12
                    elif "am" in start_phrase.lower() and hour == 12: hour = 0

                final_start = target_date.replace(hour=hour, minute=minute, second=0, microsecond=0)
                
                appt.Start = final_start.strftime("%Y-%m-%d %H:%M")
                appt.Duration = int(duration)
                appt.Subject = str(subject).title()
                appt.MeetingStatus = 1 # olMeeting
                
                # --- THE FIX: Add a default location to bypass the warning popup ---
                appt.Location = "Microsoft Teams" 

                for attendee in attendees:
                    appt.Recipients.Add(attendee)
                
                # We still must Display() to wake up the Teams Ribbon UI
                appt.Display()
                time.sleep(0.8) # Wait for window ribbon to load
                
                try:
                    # 1. Inject the Teams link via Ribbon macro
                    appt.GetInspector.CommandBars.ExecuteMso("AddOnlineMeeting")
                    print(f"[ACTION][MEETING] Teams link successfully requested for '{subject}'.")
                    
                    # 2. THE FIX: Increase wait time to 2.5 seconds. 
                    # This guarantees the Teams cloud finishes pasting the link before we proceed.
                    time.sleep(20.5)
                    
                    # 3. THE FIX: Force Outlook to save the newly injected link to memory
                    appt.Save()
                    
                    # 4. ACTUALLY SEND THE INVITE 
                    appt.Send()
                    print(f"[ACTION][MEETING] Invite sent successfully to {attendees}!")
                    
                except Exception as ribbon_err:
                    print(f"[ACTION][MEETING][WARN] Ribbon element slow or missing: {ribbon_err}")
            finally:
                pythoncom.CoUninitialize()

        print(f"[Short-Circuit Executing] Spawning Outlook background thread...")
        threading.Thread(target=run_outlook, daemon=True).start()

    def _execute_json_action(self, action_obj):
        action = action_obj.get("action", "")
        target = (action_obj.get("target") or "").strip()
        value = action_obj.get("value", "")

        if action == "open":
            import AppOpener,pywhatkit, openpyxl, docx, pptx, pyautogui
            if not target: return False
            app_name = AppOpener.identify_app(target)
            if app_name: return self._run_tool_bridge_action({"action": "launch_application", "app": app_name})

        if action == "close":
            return self._run_tool_bridge_action({"action": "close_application", "app": target})

        if action == "search_web":
            if not target: return False
            search_url = f"https://duckduckgo.com/?q={quote_plus(target)}"
            webbrowser.open(search_url)
            print(f"[ACTION][WEB] Searching web for: {target}")
            return True

        if action == "open_website":
            if not target: return False
            url = target if target.startswith(("http://", "https://")) else "https://" + target
            webbrowser.open(url)
            print(f"[ACTION][WEB] Opened {url}")
            return True

        if action == "volume":
            import pyautogui 
            value_text = f"{target} {value}".lower()
            if "up" in value_text: pyautogui.press("volumeup")
            elif "down" in value_text: pyautogui.press("volumedown")
            elif "mute" in value_text or "unmute" in value_text: pyautogui.press("volumemute")
            else: return False
            return True

        if action == "write_note":
            import pyautogui 
            content = target or str(value)
            if not content: return False
            self._open_text_editor()
            time.sleep(1.0)
            pyautogui.write(content, interval=0.05)
            print(f"[ACTION] Writing to editor: {content}")
            return True

        if action == "play":
            if not target: return False
            if pywhatkit is not None:
                pywhatkit.playonyt(target)
                print(f"[ACTION] Playing on YouTube: {target}")
                return True
            else: return False

        return False

    def execute_from_assistant(self, text):
        action_obj = self._extract_json_action(text)
        if not action_obj:
            print("[ACTION][SAFE] Assistant output has no valid JSON action. Ignored.")
            return False
        return self._execute_json_action(action_obj)

    def execute(self, text):
        if not text: return
        print(f"[ACTION][DEBUG] Executing: '{text}'")
        try:
            raw_text = text.strip()
            
            if raw_text.startswith("```json"):
                raw_text = raw_text.replace("```json", "").replace("```", "").strip()

            try:
                potential_json = json.loads(raw_text)
                if isinstance(potential_json, dict) and "action" in potential_json:
                    print(json.dumps(run_tool_action(potential_json), ensure_ascii=False))
                    return
            except json.JSONDecodeError:
                pass

            text = raw_text.lower()
            structured_action = self._extract_json_action(raw_text)
            if structured_action: return self._execute_json_action(structured_action)

            if self._handle_extended_tool_commands(raw_text): return
            if self.excel.handle(raw_text): return
            if text.startswith("word") and self.word.handle(raw_text): return
            if text.startswith("powerpoint") and self.powerpoint.handle(raw_text): return

            if text == "office help":
                self.excel.handle("excel help")
                self.word.handle("word help")
                self.powerpoint.handle("powerpoint help")
                
            if text.startswith("schedule") or "meeting" in text:
                emails = re.findall(r"[\w\.-]+@[\w\.-]+", raw_text)
                if emails:
                    # 1. Extract Duration (default to 30 mins)
                    duration_match = re.search(r"(\d+)\s*(?:minute|min|hour)", text)
                    duration = int(duration_match.group(1)) if duration_match else 30
                    if "hour" in text and duration < 5: 
                        duration *= 60
                    
                    # 2. Extract Subject Line
                    subject = "Sync Meeting"
                    subject_match = re.search(r"(?:schedule|create)\s+(?:an?\s+)?(?:\d+-minute\s+)?(.+?\s+meeting)", raw_text, re.IGNORECASE)
                    if subject_match:
                        subject = subject_match.group(1).strip()
                    elif "fyp" in text:
                        subject = "FYP Sync Meeting"

                    # 3. Extract Human Time Phrase
                    start_phrase = "tomorrow at 10:00 AM"
                    time_match = re.search(r"\b(tomorrow|today|monday|tuesday|wednesday|thursday|friday)\s*(?:at\s*)?(\d+(?::\d+)?\s*(?:am|pm)?)", text, re.IGNORECASE)
                    if time_match:
                        start_phrase = f"{time_match.group(1)} at {time_match.group(2)}"

                    # 4. Trigger our new thread-isolated background function
                    self._run_meeting_automation(subject, start_phrase, duration, emails)    
                return

            if match := self.CMD_WEB_RESEARCH_RX.fullmatch(raw_text): return self._research_and_store_web_data(match.group(1))
            
            if match := self.CMD_OPEN_WEB_RX.fullmatch(raw_text):
                url = match.group(1).strip()
                url = url if url.startswith(("http://", "https://")) else "https://" + url
                webbrowser.open(url)
                print(f"[ACTION][WEB] Opened {url}")
                return

            if match := self.CMD_SEARCH_WEB_RX.fullmatch(raw_text):
                query = match.group(1).strip()
                search_url = f"[https://duckduckgo.com/?q=](https://duckduckgo.com/?q=){quote_plus(query)}"
                webbrowser.open(search_url)
                print(f"[ACTION][WEB] Searching web for: {query}")
                return

            if self.CMD_COPY_TEXT_RX.fullmatch(text):
                pyautogui.hotkey("ctrl", "c")
                time.sleep(0.2)
                if CLIPBOARD_AVAILABLE:
                    snippet = pyperclip.paste().strip()
                    preview = snippet[:120] + ("..." if len(snippet) > 120 else "")
                    print(f"[ACTION][CLIPBOARD] Copied: {preview}")
                else: print("[ACTION][CLIPBOARD] Copied selection (clipboard preview unavailable).")
                return

            if self.CMD_PASTE_TEXT_RX.fullmatch(text):
                pyautogui.hotkey("ctrl", "v")
                print("[ACTION][CLIPBOARD] Pasted clipboard into active software.")
                return

            if match := self.CMD_SAVE_CLIP_RX.fullmatch(raw_text):
                if not CLIPBOARD_AVAILABLE:
                    print("[ACTION][RAD] pyperclip not installed. Run: pip install pyperclip")
                    return
                if not self.db or not hasattr(self.db, "add_rad_data_if_new"):
                    print("[ACTION][RAD] Database connection unavailable for clipboard save.")
                    return
                key_name = match.group(1).strip().strip('"').strip("'")
                clip_text = pyperclip.paste().strip()
                if not clip_text:
                    print("[ACTION][RAD] Clipboard is empty.")
                    return
                if self.db.add_rad_data_if_new("clipboard_note", key_name, clip_text, 0.84): print(f"[ACTION][RAD] Clipboard saved under key '{key_name}'.")
                else: print(f"[ACTION][RAD] Duplicate clipboard note ignored for key '{key_name}'.")
                return

            if self.CMD_SYS_CHECK_RX.fullmatch(text): return self._print_system_check()
            if self.CMD_MALWARE_SCAN_RX.fullmatch(text):
                threading.Thread(target=self._run_quick_malware_scan, daemon=True).start()
                return

            if "scan apps" in text or "update apps" in text:
                print("[ACTION] Scanning for new apps...")
                if APPOPENER_AVAILABLE and give_appnames: threading.Thread(target=give_appnames, daemon=True).start()
                else: print("[ACTION] App scanning is only available when AppOpener is installed.")
                return

            if text.startswith("play"):
                video_topic = text.replace("play", "").replace("please", "").strip()
                print(f"[ACTION] Playing on YouTube: {video_topic}")
                try:
                    import pywhatkit
                    pywhatkit.playonyt(video_topic)
                except Exception as e:
                    print(f"[ERROR] pywhatkit failed, falling back to direct browser search query: {e}")
                    search_url = f"[https://www.youtube.com/results?search_query=](https://www.youtube.com/results?search_query=){quote_plus(video_topic)}"
                    webbrowser.open(search_url)
                return

            write_triggers = ["write ", "note ", "type ", "take a note "]
            triggered_word = next((w for w in write_triggers if text.startswith(w)), None)
            if triggered_word:
                content = text.replace(triggered_word, "").strip()
                print(f"[ACTION] Writing to editor: {content}")
                self._open_text_editor()
                time.sleep(1.0)
                pyautogui.write(content, interval=0.05)
                return

            if "volume up" in text: return pyautogui.press('volumeup')
            if "volume down" in text: return pyautogui.press('volumedown')
            if "mute" in text or "unmute" in text: return pyautogui.press('volumemute')

            if text.startswith("open ") or text.startswith("launch "):
                raw_name = re.sub(r'^(open|launch)\s+', '', text).strip()
                app_name = raw_name.replace("please", "").replace("now", "").strip()
                app_name = re.sub(r'[^\w\s]', '', app_name)

                if self._looks_like_file_hint(raw_name):
                    self._run_tool_bridge_action({"action": "open_file", "path": raw_name, "resolve_by_hint": True, "include_content": True})
                    return
                return self._run_tool_bridge_action({"action": "launch_application", "app": app_name})

            if text.startswith("close "):
                app_name = text.replace("close ", "").replace("please", "").strip()
                app_name = re.sub(r'[^\w\s]', '', app_name)
                return self._run_tool_bridge_action({"action": "close_application", "app": app_name})

            if self._looks_like_action_command(raw_text):
                run_tool_action({"action": "run_command", "command": raw_text})

        except Exception as global_err:
            print(f"[Critical Action Failure Guard] Automated recovery handled anomaly: {global_err}")


# =============================================================================
# 7. CLI ENTRY POINT
# =============================================================================

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run OS tools or automation action command")
    parser.add_argument("--action-json", default="", help="Tool action payload as JSON object")
    parser.add_argument("--text", default="", help="Action command text (natural language)")
    return parser.parse_args()

def _read_payload(args: argparse.Namespace) -> str:
    from_arg = str(getattr(args, "action_json", "") or "").strip()
    if from_arg: return from_arg
    try:
        if sys.stdin is not None and not sys.stdin.closed:
            return sys.stdin.read().strip()
    except Exception: pass
    return ""

def _emit(result: dict) -> None:
    print(json.dumps(result, ensure_ascii=True))

def main() -> int:
    args = _parse_args()
    text_arg = str(getattr(args, "text", "") or "").strip()
    json_arg = _read_payload(args)

    if text_arg:
        try:
            handler = ActionHandler()
            output = handler.execute_and_collect(text_arg)
            if output: print(output)
            return 0
        except Exception as exc:
            print(f"[ACTION RUNNER ERROR] {exc}")
            return 1

    if json_arg:
        try: parsed = json.loads(json_arg)
        except Exception as exc:
            _emit({"success": False, "message": "Tool action payload was not valid JSON.", "error": str(exc)})
            return 1

        if not isinstance(parsed, dict):
            _emit({"success": False, "message": "Tool action payload must be a JSON object.", "error": "invalid_action_payload"})
            return 1

        try: result = run_tool_action(parsed)
        except Exception as exc:
            _emit({"success": False, "message": "Tool action runner crashed while executing payload.", "error": str(exc)})
            return 1

        if not isinstance(result, dict):
            result = {"success": False, "message": "Tool action returned an invalid result payload.", "error": str(result)}

        _emit(result)
        return 0

    _emit({"success": False, "message": "No input provided. Use --action-json or --text.", "error": "empty_input"})
    return 1

if __name__ == "__main__":
    sys.exit(main())